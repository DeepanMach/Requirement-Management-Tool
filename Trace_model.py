# req_trace_app.py
# Integrated PyQt6 Requirement Extractor + Traceability matcher
# Author: Generated for Deepan

import os, re, threading, webbrowser
from typing import List
import pandas as pd
from PyQt6.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QLabel,
    QTextEdit, QFileDialog, QMessageBox, QTabWidget, QTableWidget, QTableWidgetItem,
    QHeaderView, QLineEdit, QComboBox
)
from PyQt6.QtGui import QTextCharFormat, QColor, QTextCursor
from PyPDF2 import PdfReader
from docx import Document
import openpyxl

# ---------------- Helpers (from your prior code) ----------------

def clean_id(text: str) -> str:
    text = text.strip("[]()").rstrip(",:;.")
    text = re.sub(r"\s+", "", text)
    text = re.sub(r"([A-Za-z0-9\-_]+-)(\d{5,})(\.\d+)?$",
                  lambda m: m.group(1) + m.group(2)[:4], text)
    text = re.sub(r"([A-Za-z0-9\-_]+-\d{1,4})\.\d{1,2}$", r"\1", text)
    return text


def build_regex_from_prefix(prefix: str) -> str:
    prefix = prefix.rstrip("-_")
    prefix = re.escape(prefix)
    return rf"{prefix}[A-Za-z0-9\-_]*[-_]\d+(?:\.\d+)?"


def build_regexes_from_input(user_input: str, pdf_mode: bool = False):
    prefixes = [p.strip() for p in user_input.replace(",", "\n").splitlines() if p.strip()]
    regex_map = {}
    for p in prefixes:
        simple_pat = build_regex_from_prefix(p)
        if pdf_mode:
            simple_pat = simple_pat.replace(
                r"\d+(?:\.\d+)?",
                r"\d+(?:\s*\d+)*(?:\.\d+)?(?=[^\dA-Za-z]|$)"
            )
        full_pat = rf"(?<![A-Za-z0-9]){simple_pat}(?![A-Za-z0-9])"
        regex_map[p] = re.compile(full_pat)
    return regex_map


def normalize_pdf_text(text: str) -> str:
    text = re.sub(r"[‐‒–—−]", "-", text)
    text = re.sub(r"\s*-\s*", "-", text)
    text = re.sub(r"\s*_\s*", "_", text)
    return text


def extract_text_from_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        chunks = []
        with open(path, "rb") as f:
            reader = PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text() or ""
                page_text = normalize_pdf_text(page_text)
                chunks.append(page_text)
        return "\n".join(chunks)

    elif ext == ".docx":
        doc = Document(path)
        chunks = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    chunks.append(cell.text)
        return "\n".join(chunks)

    elif ext in [".xlsx", ".xls"]:
        wb = openpyxl.load_workbook(path, data_only=True)
        chunks = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            for row in ws.iter_rows(values_only=True):
                for cell in row:
                    if cell:
                        chunks.append(str(cell))
        return "\n".join(chunks)

    elif ext in [".txt", ".c", ".h", ".s"]:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    return ""

# ---------------- New: extract full Requirement blocks ----------------

def extract_full_requirements(text: str) -> List[dict]:
    # captures patterns like 'Requirement ID: H398-SRS-GWY-FNC-461' (case insensitive)
    pattern = re.compile(r"Requirement\s*ID\s*:\s*([A-Z0-9\-_.]+)", re.IGNORECASE)
    matches = list(pattern.finditer(text))
    results = []
    for i, m in enumerate(matches):
        req_id = m.group(1).strip()
        start = m.end()
        end = matches[i+1].start() if i+1 < len(matches) else len(text)
        block = text[start:end].strip()
        # If block is empty, try to take the remainder of current paragraph
        if not block:
            # fallback: take next 400 chars
            block = text[start:start+400].strip()
        tables = re.findall(r"Table\s+\d+", block, flags=re.IGNORECASE)
        figures = re.findall(r"Figure\s+\d+", block, flags=re.IGNORECASE)
        results.append({
            "Requirement ID": clean_id(req_id),
            "Requirement Text": block,
            "Referenced Tables": ", ".join(sorted(set(tables))),
            "Referenced Figures": ", ".join(sorted(set(figures)))
        })
    return results

# ---------------- Trace matching helpers ----------------

try:
    from sentence_transformers import SentenceTransformer
    from sklearn.metrics.pairwise import cosine_similarity
    EMB_MODEL = SentenceTransformer('all-MiniLM-L6-v2')
except Exception:
    EMB_MODEL = None


def compute_trace_matrix(sys_reqs: List[dict], sw_reqs: List[dict], threshold: float = 0.75):
    """Return list of trace dicts between system and software requirements."""
    traces = []
    # build explicit map for quick ID reference
    sys_by_id = {r['Requirement ID']: r for r in sys_reqs}

    # direct ID reference: if software text contains system ID
    for s in sw_reqs:
        for sid, sysr in sys_by_id.items():
            if sid in s['Requirement Text']:
                traces.append({
                    'System ID': sid,
                    'Software ID': s['Requirement ID'],
                    'Method': 'Direct Reference',
                    'Confidence': 0.99,
                    'Evidence': f"Software text mentions {sid}"
                })

    # table-based linking
    for s in sw_reqs:
        for sysr in sys_reqs:
            if sysr['Referenced Tables'] and s['Referenced Tables']:
                common = set([t.lower() for t in sysr['Referenced Tables'].split(',') if t.strip()]) & \
                         set([t.lower() for t in s['Referenced Tables'].split(',') if t.strip()])
                if common:
                    traces.append({
                        'System ID': sysr['Requirement ID'],
                        'Software ID': s['Requirement ID'],
                        'Method': 'Table Match',
                        'Confidence': 0.95,
                        'Evidence': f"Shared tables: {', '.join(common)}"
                    })

    # semantic similarity using embeddings (if model available)
    if EMB_MODEL and sys_reqs and sw_reqs:
        sys_texts = [r['Requirement Text'] for r in sys_reqs]
        sw_texts = [r['Requirement Text'] for r in sw_reqs]
        sys_emb = EMB_MODEL.encode(sys_texts, show_progress_bar=False)
        sw_emb = EMB_MODEL.encode(sw_texts, show_progress_bar=False)
        sim = cosine_similarity(sw_emb, sys_emb)
        for i, s in enumerate(sw_reqs):
            for j, sysr in enumerate(sys_reqs):
                score = float(sim[i, j])
                if score >= threshold:
                    traces.append({
                        'System ID': sysr['Requirement ID'],
                        'Software ID': s['Requirement ID'],
                        'Method': 'Semantic Match',
                        'Confidence': round(score, 3),
                        'Evidence': 'semantic_similarity'
                    })
    return traces

# ---------------- Main PyQt App ----------------
class ReqTraceApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('Req Extractor + Trace Tool')
        self.resize(1200, 800)
        self.file_path = None
        self.regexes_map = {}
        self.grouped_results = None
        self.pdf_mode = False
        self.sys_reqs = []
        self.sw_reqs = []
        self.traces = []

        layout = QVBoxLayout()
        self.tabs = QTabWidget()
        layout.addWidget(self.tabs)

        # Extractor Tab (adapted)
        self.tab_extract = QWidget()
        exlayout = QVBoxLayout(self.tab_extract)
        exlayout.addWidget(QLabel('Prefixes (comma/line separated):'))
        self.prefix_input = QTextEdit(); self.prefix_input.setFixedHeight(80)
        self.prefix_input.setPlaceholderText('Examples:\nCAP-SRS-\nH398-SRS-GWY-\nDU_SYS_')
        exlayout.addWidget(self.prefix_input)

        btn_row = QHBoxLayout()
        self.btn_file = QPushButton('Select File'); self.btn_folder = QPushButton('Select Folder')
        self.btn_preview = QPushButton('Preview Regexes'); self.btn_highlight = QPushButton('Highlight Matches')
        btn_row.addWidget(self.btn_file); btn_row.addWidget(self.btn_folder); btn_row.addWidget(self.btn_preview); btn_row.addWidget(self.btn_highlight)
        exlayout.addLayout(btn_row)

        exlayout.addWidget(QLabel('Regex Preview:'))
        self.regex_preview = QTextEdit(); self.regex_preview.setReadOnly(True); self.regex_preview.setFixedHeight(100)
        exlayout.addWidget(self.regex_preview)

        exlayout.addWidget(QLabel('Document Text:'))
        self.text_edit = QTextEdit(); exlayout.addWidget(self.text_edit)

        # extract buttons
        ex_buttons = QHBoxLayout()
        self.btn_extract_full = QPushButton('Extract Full Requirements')
        self.btn_extract_ids = QPushButton('Extract IDs (legacy)')
        self.btn_save = QPushButton('Save Extracted')
        ex_buttons.addWidget(self.btn_extract_full); ex_buttons.addWidget(self.btn_extract_ids); ex_buttons.addWidget(self.btn_save)
        exlayout.addLayout(ex_buttons)

        self.table = QTableWidget(); exlayout.addWidget(self.table)
        self.status_ex = QLabel('No data extracted.'); exlayout.addWidget(self.status_ex)
        self.tab_extract.setLayout(exlayout)
        self.tabs.addTab(self.tab_extract, 'Extractor')

        # Trace Tab
        self.tab_trace = QWidget()
        trlayout = QVBoxLayout(self.tab_trace)
        # controls for selecting system and software source (either doc/pdf or extracted excel)
        sel_row = QHBoxLayout()
        sel_row.addWidget(QLabel('System Source:'))
        self.sys_combo = QComboBox(); self.sys_combo.addItems(['Select file','Doc/PDF/DOCX','Extractor Excel'])
        sel_row.addWidget(self.sys_combo)
        self.btn_sys = QPushButton('Browse System'); sel_row.addWidget(self.btn_sys)
        sel_row.addSpacing(20)
        sel_row.addWidget(QLabel('Software Source:'))
        self.sw_combo = QComboBox(); self.sw_combo.addItems(['Select file','Doc/PDF/DOCX','Extractor Excel'])
        sel_row.addWidget(self.sw_combo)
        self.btn_sw = QPushButton('Browse Software'); sel_row.addWidget(self.btn_sw)
        trlayout.addLayout(sel_row)

        run_row = QHBoxLayout()
        self.btn_run_trace = QPushButton('Run Trace Matching')
        self.btn_export_trace = QPushButton('Export Trace Matrix')
        run_row.addWidget(self.btn_run_trace); run_row.addWidget(self.btn_export_trace)
        trlayout.addLayout(run_row)

        trlayout.addWidget(QLabel('Trace Results:'))
        self.trace_table = QTableWidget(); trlayout.addWidget(self.trace_table)
        self.status_trace = QLabel('No traces computed yet.'); trlayout.addWidget(self.status_trace)
        self.tab_trace.setLayout(trlayout)
        self.tabs.addTab(self.tab_trace, 'Trace')

        self.setLayout(layout)

        # connections
        self.btn_file.clicked.connect(self.select_file)
        self.btn_folder.clicked.connect(self.select_folder)
        self.btn_preview.clicked.connect(self.preview_regexes)
        self.btn_highlight.clicked.connect(self.highlight_matches)
        self.btn_extract_full.clicked.connect(self.extract_full_from_editor)
        self.btn_extract_ids.clicked.connect(self.extract_ids_legacy)
        self.btn_save.clicked.connect(self.save_extracted)

        self.btn_sys.clicked.connect(lambda: self.browse_trace_source(side='sys'))
        self.btn_sw.clicked.connect(lambda: self.browse_trace_source(side='sw'))
        self.btn_run_trace.clicked.connect(self.start_trace_matching)
        self.btn_export_trace.clicked.connect(self.export_traces)

    # ---------------- File selection / folder ----------------
    def select_file(self):
        path, _ = QFileDialog.getOpenFileName(self, 'Select File', '',
            'All supported (*.pdf *.docx *.doc *.xlsx *.xls *.txt *.c *.h *.s);;PDF (*.pdf);;Word (*.docx *.doc);;Excel (*.xlsx *.xls);;Text (*.txt)')
        if not path:
            return
        self.file_path = path
        self.pdf_mode = path.lower().endswith('.pdf')
        text = extract_text_from_file(path)
        if not text.strip():
            QMessageBox.warning(self, 'Error', 'Could not extract text from file.')
            return
        user_input = self.prefix_input.toPlainText()
        if user_input.strip():
            self.regexes_map = build_regexes_from_input(user_input, pdf_mode=self.pdf_mode)
            preview = '\n'.join([f"{p} → {r.pattern} {'[PDF-mode]' if self.pdf_mode else ''}" for p, r in self.regexes_map.items()])
            self.regex_preview.setPlainText(preview)
        self.text_edit.setPlainText(text)
        self.status_ex.setText(f'Loaded file: {os.path.basename(path)}')

    def select_folder(self):
        folder = QFileDialog.getExistingDirectory(self, 'Select Folder', '')
        if not folder:
            return
        self.file_path = folder
        self.pdf_mode = False
        supported = ('.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt', '.c', '.h', '.s')
        texts = []
        file_count = 0
        for root, _, files in os.walk(folder):
            for fname in files:
                if fname.lower().endswith(supported):
                    fpath = os.path.join(root, fname)
                    try:
                        file_text = extract_text_from_file(fpath)
                        if file_text.strip():
                            texts.append(f"--- {fname} ---\n{file_text}")
                            file_count += 1
                    except Exception as e:
                        print(f"[WARN] Could not read {fpath}: {e}")
        if not texts:
            QMessageBox.warning(self, 'Error', 'No supported files found or could not extract text.')
            return
        combined = '\n\n'.join(texts)
        user_input = self.prefix_input.toPlainText()
        if user_input.strip():
            self.regexes_map = build_regexes_from_input(user_input, pdf_mode=False)
            preview = '\n'.join([f"{p} → {r.pattern}" for p, r in self.regexes_map.items()])
            self.regex_preview.setPlainText(preview)
        self.text_edit.setPlainText(combined)
        self.status_ex.setText(f'Loaded folder: {os.path.basename(folder)} ({file_count} files)')

    def preview_regexes(self):
        user_input = self.prefix_input.toPlainText()
        if not user_input.strip():
            QMessageBox.warning(self, 'Error', 'Enter at least one prefix.')
            return
        self.regexes_map = build_regexes_from_input(user_input, pdf_mode=self.pdf_mode)
        preview = '\n'.join([f"{p} → {r.pattern} {'[PDF-mode]' if self.pdf_mode else ''}" for p, r in self.regexes_map.items()])
        self.regex_preview.setPlainText(preview)
        self.status_ex.setText('Regexes generated.')

    # ---------------- Highlight ----------------
    def highlight_matches(self):
        if not self.regexes_map:
            QMessageBox.warning(self, 'Error', 'Generate regexes first.')
            return
        text = self.text_edit.toPlainText()
        cursor = self.text_edit.textCursor()
        cursor.beginEditBlock()
        cursor.select(QTextCursor.SelectionType.Document)
        cursor.setCharFormat(QTextCharFormat())
        cursor.clearSelection()
        colors = ['yellow','lightgreen','lightblue','orange','pink','violet','cyan']
        all_results = {}
        for idx, (pat, regex) in enumerate(self.regexes_map.items()):
            fmt = QTextCharFormat(); fmt.setBackground(QColor(colors[idx % len(colors)]))
            matches = regex.findall(text)
            cleaned = [clean_id(m) for m in matches]
            all_results[pat] = sorted(set(cleaned))
            for match in regex.finditer(text):
                cursor.setPosition(match.start())
                cursor.movePosition(QTextCursor.MoveOperation.Right, QTextCursor.MoveMode.KeepAnchor, match.end() - match.start())
                cursor.mergeCharFormat(fmt)
        cursor.endEditBlock()
        self.grouped_results = all_results
        total = sum(len(v) for v in all_results.values())
        self.status_ex.setText(f'Highlighted {total} unique IDs across {len(all_results)} prefix(es).')

    # ---------------- Extraction: full requirements ----------------
    def extract_full_from_editor(self):
        text = self.text_edit.toPlainText()
        req_blocks = extract_full_requirements(text)
        if not req_blocks:
            QMessageBox.warning(self, 'No Requirements', "No 'Requirement ID:' patterns found in the document.")
            return
        self.grouped_results = req_blocks
        # populate table with 4 columns
        self.table.setRowCount(len(req_blocks))
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(['Requirement ID','Requirement Text','Referenced Tables','Referenced Figures'])
        for r, req in enumerate(req_blocks):
            self.table.setItem(r, 0, QTableWidgetItem(req['Requirement ID']))
            self.table.setItem(r, 1, QTableWidgetItem(req['Requirement Text']))
            self.table.setItem(r, 2, QTableWidgetItem(req['Referenced Tables']))
            self.table.setItem(r, 3, QTableWidgetItem(req['Referenced Figures']))
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.status_ex.setText(f'Extracted {len(req_blocks)} full requirements.')

    # legacy simple ID extraction
    def extract_ids_legacy(self):
        text = self.text_edit.toPlainText()
        if not self.regexes_map:
            QMessageBox.warning(self, 'Error', 'Generate regexes first (Preview Regexes).'); return
        grouped_results = {}
        for pat, regex in self.regexes_map.items():
            found = regex.findall(text)
            cleaned = [clean_id(m) for m in found]
            grouped_results[pat] = sorted(set(cleaned))
        self.grouped_results = grouped_results
        all_ids = [(i, pat) for pat, ids in grouped_results.items() for i in ids]
        self.table.setRowCount(len(all_ids)); self.table.setColumnCount(2)
        self.table.setHorizontalHeaderLabels(['Requirement ID','Prefix'])
        for row, (rid, pat) in enumerate(all_ids):
            self.table.setItem(row, 0, QTableWidgetItem(rid)); self.table.setItem(row, 1, QTableWidgetItem(pat))
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.status_ex.setText(f'Extracted {len(all_ids)} unique IDs.')

    def save_extracted(self):
        if not self.grouped_results:
            QMessageBox.warning(self,'Error','No extracted IDs to save.'); return
        path, _ = QFileDialog.getSaveFileName(self,'Save Results','','Excel files (*.xlsx)')
        if not path: return
        # if grouped_results is list of dicts (full requirements)
        if isinstance(self.grouped_results, list) and self.grouped_results and 'Requirement Text' in self.grouped_results[0]:
            df = pd.DataFrame(self.grouped_results)
            df.to_excel(path, index=False)
            QMessageBox.information(self,'Saved',f'Results saved to {path}'); self.status_ex.setText(f'Results saved to {path}'); return
        # else grouped_results is dict
        with pd.ExcelWriter(path, engine='openpyxl') as writer:
            all_rows = []
            for pat, ids in (self.grouped_results or {}).items():
                if ids:
                    df = pd.DataFrame(ids, columns=['Requirement ID'])
                    sheet = pat[:31] or 'IDs'
                    df.to_excel(writer, sheet_name=sheet, index=False)
                    for i in ids: all_rows.append((i, pat))
            if all_rows:
                df_all = pd.DataFrame(all_rows, columns=['Requirement ID','Prefix']).drop_duplicates().sort_values('Requirement ID')
                df_all.to_excel(writer, sheet_name='All_IDs', index=False)
        QMessageBox.information(self,'Saved',f'Results saved to {path}'); self.status_ex.setText(f'Results saved to {path}')

    # ---------------- Trace Tab: browse sources ----------------
    def browse_trace_source(self, side='sys'):
        combo = self.sys_combo if side=='sys' else self.sw_combo
        mode = combo.currentText()
        if mode == 'Select file':
            QMessageBox.warning(self,'Select mode','Please choose source type (Doc/PDF or Extractor Excel) from the dropdown.'); return
        if mode == 'Doc/PDF/DOCX':
            path, _ = QFileDialog.getOpenFileName(self,'Select Document','', 'Docs (*.pdf *.docx *.doc)')
            if not path: return
            text = extract_text_from_file(path)
            reqs = extract_full_requirements(text)
        else:
            path, _ = QFileDialog.getOpenFileName(self,'Select Extractor Excel','', 'Excel (*.xlsx *.xls)')
            if not path: return
            try:
                df = pd.read_excel(path)
                # expect columns: Requirement ID, Requirement Text (best) or similar
                if 'Requirement ID' in df.columns and 'Requirement Text' in df.columns:
                    reqs = df[['Requirement ID','Requirement Text','Referenced Tables','Referenced Figures']].fillna('').to_dict(orient='records')
                elif 'Requirement ID' in df.columns:
                    reqs = [{'Requirement ID':row['Requirement ID'], 'Requirement Text': row.get('Requirement Text',''), 'Referenced Tables':'', 'Referenced Figures':''} for _, row in df.iterrows()]
                else:
                    QMessageBox.warning(self,'Invalid','Excel does not appear to be extractor output.'); return
            except Exception as e:
                QMessageBox.warning(self,'Error',f'Could not read Excel: {e}'); return
        if side=='sys':
            self.sys_reqs = reqs; self.status_trace.setText(f'Loaded {len(reqs)} system requirements')
        else:
            self.sw_reqs = reqs; self.status_trace.setText(f'Loaded {len(reqs)} software requirements')

    def start_trace_matching(self):
        if not self.sys_reqs or not self.sw_reqs:
            QMessageBox.warning(self,'Missing Data','Please load both System and Software sources first.')
            return
        # run matching in background thread
        threading.Thread(target=self.run_matching_thread, daemon=True).start()

    def run_matching_thread(self):
        self.status_trace.setText('Running trace matching...')
        try:
            traces = compute_trace_matrix(self.sys_reqs, self.sw_reqs, threshold=0.75)
            if not traces:
                self.status_trace.setText('No traces found.'); QMessageBox.information(self,'No Traces','No candidate traces found.'); return
            self.traces = traces
            # populate trace_table
            cols = ['System ID','Software ID','Method','Confidence','Evidence']
            self.trace_table.setRowCount(len(traces)); self.trace_table.setColumnCount(len(cols))
            self.trace_table.setHorizontalHeaderLabels(cols)
            for r, t in enumerate(traces):
                self.trace_table.setItem(r,0,QTableWidgetItem(t['System ID']))
                self.trace_table.setItem(r,1,QTableWidgetItem(t['Software ID']))
                self.trace_table.setItem(r,2,QTableWidgetItem(t['Method']))
                self.trace_table.setItem(r,3,QTableWidgetItem(str(t['Confidence'])))
                self.trace_table.setItem(r,4,QTableWidgetItem(t.get('Evidence','')))
            self.trace_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
            self.status_trace.setText(f'Found {len(traces)} candidate traces.')
        except Exception as e:
            QMessageBox.critical(self,'Error',str(e)); self.status_trace.setText('Error during trace matching')

    def export_traces(self):
        if not getattr(self,'traces',None):
            QMessageBox.warning(self,'No Data','No trace results to export.'); return
        path, _ = QFileDialog.getSaveFileName(self,'Export Trace Matrix','trace_matrix.xlsx','Excel (*.xlsx *.xls)')
        if not path: return
        df = pd.DataFrame(self.traces)
        df.to_excel(path, index=False)
        QMessageBox.information(self,'Saved',f'Trace matrix exported to {path}')

# ---------------- Run ----------------
if __name__ == '__main__':
    import sys
    app = QApplication(sys.argv)
    win = ReqTraceApp()
    win.show()
    sys.exit(app.exec())
