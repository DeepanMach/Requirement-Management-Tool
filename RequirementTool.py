import sys, os, re
import pandas as pd
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QTableWidget, QTableWidgetItem, QPushButton, QFileDialog,
    QTextEdit, QTreeWidget, QTreeWidgetItem, QMessageBox,
    QHeaderView, QAbstractItemView, QMenu, QInputDialog
)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QAction, QShortcut, QKeySequence
from openpyxl import Workbook

# ============================================================
# Main Application
# ============================================================
class ExcelToWordApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Mach Requirement Management Tool")
        self.resize(1300, 800)

        self.df = pd.DataFrame()
        self.filtered_df = pd.DataFrame()

        # === Layout Setup ===
        main = QWidget()
        self.setCentralWidget(main)
        self.main_layout = QVBoxLayout(main)

        # --- Buttons Row ---
        btn_layout = QHBoxLayout()
        self.load_btn = QPushButton("📂 Load Excel")
        self.load_btn.clicked.connect(self.load_excels)

        self.run_btn = QPushButton("▶️ Run (Convert to Word)")
        self.run_btn.clicked.connect(self.run_convert_preview)

        self.back_btn = QPushButton("🔙 Back to Table View")
        self.back_btn.clicked.connect(self.show_table_view)

        self.save_btn = QPushButton("💾 Save Word")
        self.save_btn.clicked.connect(self.save_word)

        self.trace_btn = QPushButton("🔗 Traceability Matrix")
        self.trace_btn.clicked.connect(self.show_trace_view)

        btn_layout.addWidget(self.load_btn)
        btn_layout.addWidget(self.run_btn)
        btn_layout.addWidget(self.back_btn)
        btn_layout.addWidget(self.save_btn)
        btn_layout.addWidget(self.trace_btn)
        self.main_layout.addLayout(btn_layout)

        # --- Split navigation and stacked area ---
        center_layout = QHBoxLayout()
        self.main_layout.addLayout(center_layout)

        # Navigation tree
        self.nav_tree = QTreeWidget()
        self.nav_tree.setHeaderLabel("Navigation")
        self.nav_tree.itemClicked.connect(self.on_nav_item_clicked)
        center_layout.addWidget(self.nav_tree, 2)

        # --- Stacked area (3 views) ---
        from PyQt6.QtWidgets import QStackedWidget
        self.view_stack = QStackedWidget()
        center_layout.addWidget(self.view_stack, 8)

        # Table view
        self.table = QTableWidget()
        self.table.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked)
        self.table.itemChanged.connect(self.on_cell_changed)
        self.table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.table.customContextMenuRequested.connect(self.table_context_menu)
        self.view_stack.addWidget(self.table)

        # Word preview
        self.word_preview = QTextEdit()
        self.word_preview.setReadOnly(True)
        self.view_stack.addWidget(self.word_preview)

        # Traceability matrix view placeholder (added in Part 2)
        from PyQt6.QtWidgets import QLabel
        self.trace_placeholder = QLabel("Traceability Matrix view will load...")
        self.view_stack.addWidget(self.trace_placeholder)

        # Console log
        from PyQt6.QtWidgets import QTextBrowser
        self.console = QTextBrowser()
        self.console.setFixedHeight(120)
        self.main_layout.addWidget(self.console)

        # Undo shortcut
        QShortcut(QKeySequence("Ctrl+Z"), self, activated=self.undo_last)

        # Vars
        self._undo_stack = []
        self.console.append("✅ Application Ready.\n")

    # ============================================================
    # Excel Loading and Table Population
    # ============================================================
    def load_excels(self):
        files, _ = QFileDialog.getOpenFileNames(self, "Select Excel Files", "", "Excel Files (*.xlsx *.xls)")
        if not files:
            return
        all_dfs = []
        for f in files:
            try:
                xls = pd.ExcelFile(f)
                for s in xls.sheet_names:
                    df = pd.read_excel(f, sheet_name=s).fillna("")
                    df["SourceFile"] = os.path.basename(f)
                    df["SheetName"] = s
                    all_dfs.append(df)
            except Exception as e:
                self.console.append(f"❌ Failed to read {f}: {e}\n")

        if not all_dfs:
            return

        combined = pd.concat(all_dfs, ignore_index=True)
        combined = combined.drop_duplicates(ignore_index=True)
        self.df = combined.copy()
        self.filtered_df = self.df.copy()

        self.apply_numbering_once_by_heading()
        self.populate_table()
        self.console.append("✅ Excel files loaded and numbering applied.\n")

    # ------------------------------------------------------------
    def apply_numbering_once_by_heading(self):
        if "Object Type" not in self.df.columns:
            return
        num_h1 = num_h2 = num_h3 = 0
        section_numbers = []
        for _, row in self.df.iterrows():
            obj_type = str(row.get("Object Type", "")).strip().lower()
            if obj_type == "heading 1":
                num_h1 += 1; num_h2 = num_h3 = 0
                section_numbers.append(f"{num_h1}.")
            elif obj_type == "heading 2":
                if num_h1 == 0: num_h1 = 1
                num_h2 += 1; num_h3 = 0
                section_numbers.append(f"{num_h1}.{num_h2}")
            elif obj_type == "heading 3":
                if num_h1 == 0: num_h1 = 1
                if num_h2 == 0: num_h2 = 1
                num_h3 += 1
                section_numbers.append(f"{num_h1}.{num_h2}.{num_h3}")
            else:
                section_numbers.append("")
        self.df.insert(0, "Section Number", section_numbers)

    
    def populate_table(self):
        """Populate the editable table view safely and style headings."""
        if self.df.empty:
            return

        # Hide "Object Type" from visible columns but keep for logic
        visible_cols = [c for c in self.df.columns if c.lower().strip() != "object type"]
        self.table.clear()
        self._loading = True
        self.table.setRowCount(len(self.df))
        self.table.setColumnCount(len(visible_cols))
        self.table.setHorizontalHeaderLabels(visible_cols)
        self.table.horizontalHeader().setStretchLastSection(True)

        for r in range(len(self.df)):
            for ci, col in enumerate(visible_cols):
                val = str(self.df.iloc[r][col])
                item = QTableWidgetItem(val)

                # Bold styling for headings
                if "heading" in str(self.df.iloc[r].get("Object Type", "")).lower():
                    f = item.font()
                    f.setBold(True)
                    item.setFont(f)

                self.table.setItem(r, ci, item)

        self.table.resizeColumnsToContents()
        self._loading = False
        self.view_stack.setCurrentIndex(0)
        self.console.append("📊 Table populated with Excel-style formatting.\n")
        self.populate_navigation()
    # ------------------------------------------------------------
    def populate_navigation(self):
        self.nav_tree.clear()
        if self.df.empty or "Object Type" not in self.df.columns:
            return

        seen = set()
        parent_h1 = None; parent_h2 = None
        for _, row in self.df.iterrows():
            obj = str(row.get("Object Type", "")).lower()
            text = str(row.get("Object Text", "")).strip()
            num = str(row.get("Section Number", "")).strip()
            if not text: continue
            label = f"{num} {text}".strip()

            if obj == "heading 1":
                if label.lower() in seen: continue
                parent_h1 = QTreeWidgetItem([label]); self.nav_tree.addTopLevelItem(parent_h1)
                seen.add(label.lower()); parent_h2 = None
            elif obj == "heading 2":
                if not parent_h1: continue
                child = QTreeWidgetItem([label]); parent_h1.addChild(child)
                parent_h2 = child
            elif obj == "heading 3":
                if parent_h2: parent_h2.addChild(QTreeWidgetItem([label]))
        self.nav_tree.expandAll()

    # ------------------------------------------------------------
    def on_nav_item_clicked(self, item):
        """Scroll Word preview to the selected heading in navigation."""
        try:
            target = item.text(0).split(" ", 1)[-1].strip()
            from PyQt6.QtGui import QTextCursor

            # Move cursor to start
            self.word_preview.moveCursor(QTextCursor.MoveOperation.Start)

            # Find the heading text in the HTML content
            cursor = self.word_preview.textCursor()
            found = self.word_preview.find(target)
            if found:
                self.word_preview.setTextCursor(cursor)
                self.word_preview.ensureCursorVisible()
                self.console.append(f"🧭 Navigated to: {target}\n")
            else:
                self.console.append(f"⚠️ Could not find heading: {target}\n")
        except Exception as e:
            self.console.append(f"❌ Navigation error: {e}\n")


    
    def on_cell_changed(self, item):
        """Update DataFrame when a cell is edited; safely handle column mismatch."""
        if getattr(self, "_loading", False):
            return

        row = item.row()
        col = item.column()
        new_val = item.text().strip()

        visible_cols = [self.table.horizontalHeaderItem(i).text() for i in range(self.table.columnCount())]

        # Ensure df has same columns
        for colname in visible_cols:
            if colname not in self.df.columns:
                self.df[colname] = ""

        if row >= len(self.df):
            self.console.append(f"⚠️ Invalid row index: {row}\n")
            return

        if col < len(visible_cols):
            colname = visible_cols[col]
            try:
                self.df.at[row, colname] = new_val
                self.filtered_df = self.df.copy()
                self._undo_stack.append(self.filtered_df.copy())
                # update navigation if user edits a heading
                obj_type = str(self.df.iloc[row].get("Object Type", "")).lower()
                if "heading" in obj_type:
                    self.populate_navigation()
            except Exception as e:
                self.console.append(f"❌ Edit failed: {e}\n")


    # ------------------------------------------------------------
    def table_context_menu(self, pos):
        menu = QMenu(self)
        add_row = menu.addAction("Add Row Below")
        del_row = menu.addAction("Delete Row")
        act = menu.exec(self.table.viewport().mapToGlobal(pos))
        if act == add_row:
            self.table.insertRow(self.table.currentRow()+1)
        elif act == del_row:
            self.table.removeRow(self.table.currentRow())

    # ------------------------------------------------------------
    def undo_last(self):
        if len(self._undo_stack) > 1:
            self._undo_stack.pop()
            self.filtered_df = self._undo_stack[-1].copy()
            self.populate_table()
            self.console.append("↩️ Undo applied.\n")

    # ============================================================
    # Word Conversion / Preview
    # ============================================================
    def run_convert_preview(self):
        if self.df.empty: return
        html_parts = ["<div>"]
        for _, row in self.df.iterrows():
            t = str(row.get("Object Type", "")).lower()
            num = str(row.get("Section Number", "")).strip()
            text = str(row.get("Object Text", "")).strip()
            rid = str(row.get("Requirement ID", "")).strip()

            if t == "heading 1":
                html_parts.append(f"<h1>{num} {text}</h1>")
            elif t == "heading 2":
                html_parts.append(f"<h2>{num} {text}</h2>")
            elif t == "heading 3":
                html_parts.append(f"<h3>{num} {text}</h3>")
            elif rid:
                html_parts.append(f"<b>Requirement ID:</b> {rid}<br><p>{text}</p>")
            else:
                html_parts.append(f"<p>{text}</p>")
        html_parts.append("</div>")
        self.word_preview.setHtml("\n".join(html_parts))
        self.view_stack.setCurrentIndex(1)
        self.console.append("📝 Word preview generated.\n")

    # ------------------------------------------------------------
    def show_table_view(self):
        self.view_stack.setCurrentIndex(0)
        self.console.append("📊 Returned to table view.\n")

    # ------------------------------------------------------------
    def save_word(self):
        if self.word_preview.toPlainText().strip() == "":
            QMessageBox.warning(self, "Empty", "Nothing to save!")
            return
        file, _ = QFileDialog.getSaveFileName(self, "Save Word File", "", "Word Document (*.docx)")
        if not file: return
        if not file.endswith(".docx"): file += ".docx"
        from docx import Document
        doc = Document()
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(self.word_preview.toHtml(), "html.parser")
        for tag in soup.find_all(["h1","h2","h3","p","b"]):
            if tag.name == "h1": doc.add_heading(tag.get_text(), level=1)
            elif tag.name == "h2": doc.add_heading(tag.get_text(), level=2)
            elif tag.name == "h3": doc.add_heading(tag.get_text(), level=3)
            elif tag.name == "b": doc.add_paragraph(tag.get_text(), style="Strong")
            elif tag.name == "p": doc.add_paragraph(tag.get_text())
        doc.save(file)
        self.console.append(f"💾 Word file saved: {file}\n")

    # ------------------------------------------------------------
    def show_trace_view(self):
        """Switch to traceability matrix stacked view (loaded in Part 2)."""
        try:
            self.trace_matrix_view.load_data(self.df)
            self.view_stack.setCurrentIndex(2)
            self.console.append("🔗 Switched to Traceability Matrix view.\n")
        except Exception as e:
            QMessageBox.critical(self, "Trace View Error", str(e))
            self.console.append(f"❌ Trace view error: {e}\n")
    # ============================================================
# Part 2: Traceability Matrix View + App bootstrap
# ============================================================
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QTableWidget, QTableWidgetItem,
    QLabel, QComboBox, QFileDialog, QPushButton, QSizePolicy
)
from PyQt6.QtCore import QPoint

class TraceMatrixView(QWidget):
    """
    Widget to display Traceability Matrix in same stacked area.
    Supports:
      - Forward Trace (Requirement ID -> Up Trace)
      - Backward Trace (Up Trace -> Requirement ID)
      - Filter dropdowns on column headers
      - Save current view to Excel (user chooses path)
    """
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self._df = None  # backing dataframe
        self.mode = "forward"  # or "backward"

        layout = QVBoxLayout(self)
        # buttons row
        btn_row = QHBoxLayout()
        self.btn_forward = QPushButton("Forward Trace")
        self.btn_forward.clicked.connect(self.show_forward_trace)
        self.btn_backward = QPushButton("Backward Trace")
        self.btn_backward.clicked.connect(self.show_backward_trace)
        self.btn_save = QPushButton("Save to Excel")
        self.btn_save.clicked.connect(self.save_current_view)
        btn_row.addWidget(self.btn_forward)
        btn_row.addWidget(self.btn_backward)
        btn_row.addStretch()
        btn_row.addWidget(self.btn_save)
        layout.addLayout(btn_row)

        # info label
        self.info_label = QLabel("Traceability Matrix (Forward/Backward). Use header clicks to filter.")
        layout.addWidget(self.info_label)

        # table
        self.table = QTableWidget()
        self.table.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.table.horizontalHeader().setSectionsClickable(True)
        self.table.horizontalHeader().sectionClicked.connect(self._on_header_clicked)
        layout.addWidget(self.table)

        # internal: current rows and columns
        self.current_rows = []
        self.current_cols = []

        # active filters: {colname: set(values)}
        self.active_filters = {}

    def load_data(self, df: pd.DataFrame):
        """Load a DataFrame (do not modify it)."""
        self._df = df.copy() if df is not None else pd.DataFrame()
        # default to forward view on load
        self.show_forward_trace()

    # -----------------------------
    # Build forward view (Req -> UpTrace)
    # -----------------------------
    def show_forward_trace(self):
        if self._df is None or self._df.empty:
            self.table.clear()
            self.info_label.setText("No data loaded.")
            return
        cols = list(self._df.columns)
        # detect columns robustly
        req_col = next((c for c in cols if "requirement" in c.lower()), None)
        up_col = next((c for c in cols if "up trace" in c.lower() or "uptrace" in c.lower()), None)
        # fallback names
        if req_col is None:
            for tryc in ("Requirement ID", "Req ID", "ReqID"):
                if tryc in cols:
                    req_col = tryc; break
        if up_col is None:
            for c in cols:
                lc = c.lower()
                if "up" in lc and "trace" in lc:
                    up_col = c; break

        if req_col is None or up_col is None:
            QMessageBox.warning(self, "Missing Columns", "Forward Trace requires 'Requirement ID' and 'Up Trace' columns present in the loaded Excel.")
            self.table.clear()
            return

        # Build rows: (Requirement ID, Up Trace) — show Up Trace exactly as user entered
        rows = []
        for _, row in self._df.iterrows():
            rid = str(row.get(req_col, "")).strip()
            if not rid:
                continue
            up = str(row.get(up_col, "")).strip()
            rows.append((rid, up))

        self.current_cols = ["Requirement ID", "Up Trace"]
        self.current_rows = rows
        self.mode = "forward"
        self.active_filters = {}
        self._populate_table_from_current()
        self.info_label.setText(f"Forward Trace: {len(rows)} entries. (Req -> UpTrace)")

    # -----------------------------
    # Build backward view (UpTrace -> Req)
    # -----------------------------
    def show_backward_trace(self):
        if self._df is None or self._df.empty:
            self.table.clear()
            self.info_label.setText("No data loaded.")
            return
        cols = list(self._df.columns)
        req_col = next((c for c in cols if "requirement" in c.lower()), None)
        up_col = next((c for c in cols if "up trace" in c.lower() or "uptrace" in c.lower()), None)
        if req_col is None:
            for tryc in ("Requirement ID", "Req ID", "ReqID"):
                if tryc in cols:
                    req_col = tryc; break
        if up_col is None:
            for c in cols:
                lc = c.lower()
                if "up" in lc and "trace" in lc:
                    up_col = c; break
        if req_col is None or up_col is None:
            QMessageBox.warning(self, "Missing Columns", "Backward Trace requires 'Requirement ID' and 'Up Trace' columns present in the loaded Excel.")
            self.table.clear()
            return

        # Build mapping from UpTrace token -> list of Requirement IDs referencing it.
        # We'll split Up Trace cells by common separators (newline, comma, semicolon)
        mapping = {}
        for _, row in self._df.iterrows():
            rid = str(row.get(req_col, "")).strip()
            upcell = str(row.get(up_col, "")).strip()
            if not upcell or not rid:
                continue
            # split into tokens
            tokens = [t.strip() for part in str(upcell).splitlines() for t in part.split(",") for s in t.split(";") for t in [s]]
            # tokens may include blanks
            tokens = [tok for tok in tokens if tok]
            for tok in tokens:
                mapping.setdefault(tok, []).append(rid)

        # Build rows: (UpTrace ID, comma-separated Requirement IDs)
        rows = []
        for upid, reqs in mapping.items():
            rows.append((upid, ", ".join(reqs)))
        self.current_cols = ["Up Trace", "Requirement IDs"]
        self.current_rows = rows
        self.mode = "backward"
        self.active_filters = {}
        self._populate_table_from_current()
        self.info_label.setText(f"Backward Trace: {len(rows)} entries. (UpTrace -> Req)")

    # -----------------------------
    # Populate QTableWidget from current_rows/current_cols applying filters
    # -----------------------------
    def _populate_table_from_current(self):
        # apply filters
        rows = self.current_rows
        cols = self.current_cols

        # filtering if any active_filters set
        if self.active_filters:
            filt_rows = []
            for r in rows:
                include = True
                for col_idx, colname in enumerate(cols):
                    vals = self.active_filters.get(colname)
                    if vals:
                        cell_val = str(r[col_idx])
                        # if any of the selected values appears in cell_val, keep it
                        if cell_val not in vals:
                            include = False
                            break
                if include:
                    filt_rows.append(r)
            rows_to_show = filt_rows
        else:
            rows_to_show = rows

        self.table.clear()
        self.table.setColumnCount(len(cols))
        self.table.setRowCount(len(rows_to_show))
        self.table.setHorizontalHeaderLabels(cols)

        for i, row in enumerate(rows_to_show):
            for j, val in enumerate(row):
                item = QTableWidgetItem(str(val))
                item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                self.table.setItem(i, j, item)

        self.table.resizeColumnsToContents()

    # -----------------------------
    # Header clicked -> show filter menu for that column
    # -----------------------------
    def _on_header_clicked(self, logical_index):
        if logical_index < 0 or logical_index >= len(self.current_cols):
            return
        colname = self.current_cols[logical_index]
        # gather unique values for that column from current_rows
        vals = [str(r[logical_index]) for r in self.current_rows]
        unique_vals = sorted(list(dict.fromkeys(vals)))  # preserve order but unique
        # create a small popup menu with checkboxes (reusing QMenu + QWidgetAction)
        from PyQt6.QtWidgets import QMenu, QWidgetAction, QScrollArea, QWidget, QVBoxLayout, QLineEdit, QCheckBox, QPushButton
        menu = QMenu(self)
        menu.setStyleSheet("QMenu{background:#fff;border:1px solid #c0c0c0;}")

        # search box
        search = QLineEdit()
        search.setPlaceholderText("Search...")
        act_search = QWidgetAction(menu)
        act_search.setDefaultWidget(search)
        menu.addAction(act_search)
        menu.addSeparator()

        # scroll area with checkboxes
        scroll = QScrollArea()
        scroll.setMinimumWidth(320)
        scroll.setMaximumHeight(380)
        scroll.setWidgetResizable(True)
        container = QWidget()
        vbox = QVBoxLayout(container)
        vbox.setContentsMargins(6,6,6,6)
        vbox.setSpacing(4)
        checkboxes = []
        current_selected = self.active_filters.get(colname, set(unique_vals))
        for v in unique_vals:
            cb = QCheckBox(v)
            cb.setChecked(v in current_selected)
            vbox.addWidget(cb)
            checkboxes.append(cb)
        container.setLayout(vbox)
        scroll.setWidget(container)
        act_scroll = QWidgetAction(menu)
        act_scroll.setDefaultWidget(scroll)
        menu.addAction(act_scroll)

        menu.addSeparator()
        apply_btn = QPushButton("Apply")
        act_apply = QWidgetAction(menu)
        act_apply.setDefaultWidget(apply_btn)
        menu.addAction(act_apply)
        clear_act = menu.addAction("Clear Filter")
        show_all_act = menu.addAction("Show All")

        # search filter
        def on_search(text):
            s = text.lower()
            for cb in checkboxes:
                cb.setVisible(s in cb.text().lower())
        search.textChanged.connect(on_search)

        def on_apply():
            selected = [cb.text() for cb in checkboxes if cb.isChecked()]
            if len(selected) == len(checkboxes):
                # no filter (all selected)
                if colname in self.active_filters:
                    self.active_filters.pop(colname, None)
            else:
                self.active_filters[colname] = set(selected)
            self._populate_table_from_current()
            menu.close()
        apply_btn.clicked.connect(on_apply)

        def on_clear():
            if colname in self.active_filters:
                self.active_filters.pop(colname, None)
            self._populate_table_from_current()
            menu.close()
        clear_act.triggered.connect(on_clear)

        def on_show_all():
            self.active_filters = {}
            self._populate_table_from_current()
            menu.close()
        show_all_act.triggered.connect(on_show_all)

        # position menu under header section
        header = self.table.horizontalHeader()
        try:
            x = header.sectionPosition(logical_index)
        except Exception:
            x = 0
        y = header.height()
        menu.exec(self.table.mapToGlobal(QPoint(x+12, y+20)))

    # -----------------------------
    # Save current view to Excel
    # -----------------------------
    def save_current_view(self):
        if not self.current_rows or not self.current_cols:
            QMessageBox.information(self, "Nothing", "No trace data to save.")
            return
        fname, _ = QFileDialog.getSaveFileName(self, "Save Trace to Excel", "", "Excel Files (*.xlsx)")
        if not fname:
            return
        if not fname.lower().endswith(".xlsx"):
            fname += ".xlsx"
        try:
            # build pandas DataFrame and save
            df_out = pd.DataFrame(self.current_rows, columns=self.current_cols)
            df_out.to_excel(fname, index=False)
            QMessageBox.information(self, "Saved", f"Trace saved to {fname}")
            if hasattr(self.parent, "console"):
                self.parent.console.append(f"💾 Trace saved to {fname}\n")
        except Exception as e:
            QMessageBox.critical(self, "Save Error", str(e))
            if hasattr(self.parent, "console"):
                self.parent.console.append(f"❌ Failed to save trace: {e}\n")

# ------------------------------------------------------------
# Hook TraceMatrixView into main app and run
# ------------------------------------------------------------
if __name__ == "__main__":
    app = QApplication(sys.argv)
    win = ExcelToWordApp()

    # create trace matrix view and replace placeholder
    trace_view = TraceMatrixView(win)
    # remove placeholder widget (index 2) and add our trace_view in same slot
    try:
        # find placeholder (we named it trace_placeholder previously)
        if hasattr(win, "trace_placeholder"):
            # remove first then add
            win.view_stack.removeWidget(win.trace_placeholder)
    except Exception:
        pass
    # add trace_view as the third view (index 2)
    win.view_stack.addWidget(trace_view)
    win.trace_matrix_view = trace_view

    win.show()
    sys.exit(app.exec())
