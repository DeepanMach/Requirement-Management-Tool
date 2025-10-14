"""Main window for the requirement management tool."""
from __future__ import annotations

import base64
import copy
import html
import io
import json
import logging
import mimetypes
from pathlib import Path
from typing import Dict, Optional, Sequence

import pandas as pd
from PyQt6.QtCore import Qt, QTimer
from PyQt6.QtGui import QKeySequence, QShortcut, QTextCursor, QTextDocument
from PyQt6.QtWidgets import (
    QApplication,
    QFileDialog,
    QDialog,
    QDialogButtonBox,
    QDoubleSpinBox,
    QFormLayout,
    QHBoxLayout,
    QHeaderView,
    QLineEdit,
    QInputDialog,
    QMainWindow,
    QMenu,
    QMessageBox,
    QPlainTextEdit,
    QPushButton,
    QSpinBox,
    QStackedWidget,
    QTabWidget,
    QTableWidget,
    QTableWidgetItem,
    QTextBrowser,
    QTreeWidget,
    QTreeWidgetItem,
    QVBoxLayout,
    QWidget,
)

from ..data_manager import (
    OPTIONAL_METADATA_COLUMNS,
    REQUIRED_COLUMNS,
    RequirementDataError,
    RequirementDataManager,
)
from .trace_view import TraceMatrixView

LOGGER = logging.getLogger(__name__)


ASSETS_DIR = (Path(__file__).resolve().parent / "assets")
ASSETS_DIR.mkdir(exist_ok=True)
DEFAULT_HOWELL_LOGO = str(ASSETS_DIR / "howell_logo.png")
DEFAULT_MACH_LOGO = str(ASSETS_DIR / "mach_logo.png")


class MainWindow(QMainWindow):
    """Top-level window coordinating UI widgets with the data manager."""

    def __init__(self, data_manager: Optional[RequirementDataManager] = None):
        super().__init__()
        self.data_manager = data_manager or RequirementDataManager()
        self.setWindowTitle("Mach Requirement Management Tool")
        self.resize(1300, 800)
        self.setAcceptDrops(True)

        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)

        button_row = QHBoxLayout()
        self.load_btn = QPushButton('Load Excel')
        self.load_btn.clicked.connect(self.load_excels)
        self.load_word_btn = QPushButton('Load Word')
        self.load_word_btn.clicked.connect(self.load_word_documents)
        self.run_btn = QPushButton('Run (Convert to Word)')
        self.run_btn.clicked.connect(self.run_convert_preview)
        self.back_btn = QPushButton('Back to Table View')
        self.back_btn.clicked.connect(self.show_table_view)
        self.save_btn = QPushButton('Save Word')
        self.save_btn.clicked.connect(self.save_word)
        self.trace_btn = QPushButton('Traceability Matrix')
        self.trace_btn.clicked.connect(self.show_trace_view)

        for button in (
            self.load_btn,
            self.load_word_btn,
            self.run_btn,
            self.back_btn,
            self.save_btn,
            self.trace_btn,
        ):
            button_row.addWidget(button)
        main_layout.addLayout(button_row)

        attachment_row = QHBoxLayout()
        self.add_image_btn = QPushButton("Add Image")
        self.add_image_btn.clicked.connect(self.add_image_attachment)
        self.add_table_btn = QPushButton("Add Table")
        self.add_table_btn.clicked.connect(self.add_table_attachment)
        self.edit_header_btn = QPushButton("Edit Header Details")
        self.edit_header_btn.clicked.connect(self.edit_header_details)
        self.remove_tab_btn = QPushButton("Remove Tab")
        self.remove_tab_btn.clicked.connect(self.remove_current_tab)
        for button in (
            self.add_image_btn,
            self.add_table_btn,
            self.edit_header_btn,
            self.remove_tab_btn,
        ):
            attachment_row.addWidget(button)
        attachment_row.addStretch()
        main_layout.addLayout(attachment_row)


        center_row = QHBoxLayout()
        main_layout.addLayout(center_row)

        self.nav_tree = QTreeWidget()
        self.nav_tree.setHeaderLabel("Navigation")
        self.nav_tree.itemClicked.connect(self.on_nav_item_clicked)
        center_row.addWidget(self.nav_tree, 2)

        self.view_stack = QStackedWidget()
        center_row.addWidget(self.view_stack, 8)

        self.table_tabs = QTabWidget()
        self.table_tabs.currentChanged.connect(self.on_table_tab_changed)
        self.view_stack.addWidget(self.table_tabs)

        self.word_preview = QTextBrowser()
        self.view_stack.addWidget(self.word_preview)

        self.trace_view = TraceMatrixView(self)
        self.trace_view.navigate_to_requirement = self.navigate_to_requirement
        self.view_stack.addWidget(self.trace_view)

        self.console = QTextBrowser()
        self.console.setFixedHeight(120)
        main_layout.addWidget(self.console)

        QShortcut(QKeySequence("Ctrl+Z"), self, activated=self.undo_last)

        self._undo_stack: list[pd.DataFrame] = []
        self._loading = False
        self._tab_tables: Dict[str, QTableWidget] = {}
        self._tab_indices: Dict[str, list[int]] = {}
        self._tab_source_types: Dict[str, str] = {}
        self.default_header_settings = {
            "document_title": "Software Requirements Specification for Gateway Module of EDAU in UH-60X Engine Instrument System",
            "document_number": "MGT-H398-GWY-S001",
            "revision": "2.4",
            "author_name": "",
            "author_title": "(Software Engineer)",
            "reviewer_name": "",
            "reviewer_title": "(Software Engineer)",
            "qa_name": "",
            "qa_title": "(Quality Assurance Engineer)",
            "config_manager_name": "",
            "config_manager_title": "(Engineering Manager)",
            "logo_left_path": DEFAULT_HOWELL_LOGO if Path(DEFAULT_HOWELL_LOGO).exists() else "",
            "logo_right_path": DEFAULT_MACH_LOGO if Path(DEFAULT_MACH_LOGO).exists() else "",
            "preview_image_width_percent": 80,
            "export_image_width_inches": 5.5,
            "address_howell": "Howell Instruments, Inc.<br/>8945 South Freeway<br/>Fort Worth, Texas, 76140<br/>U.S.A.",
            "address_mach": "Mach Global Technologies<br/>No 42/4A, Shantipura Road, Electronic City,<br/>Phase 2, Bengaluru, Karnataka 560100, India",
            "proprietary_notice": (
                "This document and the information contained herein are the property of Howell Instruments, Inc. "
                "Any reproduction, disclosure or use thereof is prohibited except as authorized in writing by Howell Instruments, Inc. "
                "Recipient accepts the responsibility for maintaining the confidentiality of the contents of this document."
            ),
            "copyright_notice": "© 2025 Howell Instruments. All rights reserved.",
            "watermark_text": "",
        }
        self.header_profiles: Dict[str, Dict[str, object]] = {}
        self.console.append("Application Ready.")

    # ------------------------------------------------------------------
    def log_console(self, message: str) -> None:
        LOGGER.info(message)
        self.console.append(message)

    # ------------------------------------------------------------------
    def _header_key(self, source: Optional[str]) -> str:
        return source or "__default__"

    # ------------------------------------------------------------------
    def _get_header_settings(self, source: Optional[str] = None) -> Dict[str, object]:
        key = self._header_key(source or self._current_source_name())
        if key not in self.header_profiles:
            self.header_profiles[key] = copy.deepcopy(self.default_header_settings)
        return self.header_profiles[key]

    # ------------------------------------------------------------------
    def _ensure_header_profile(self, source: Optional[str]) -> None:
        self._get_header_settings(source)

    # ------------------------------------------------------------------
    def _current_table(self) -> Optional[QTableWidget]:
        widget = self.table_tabs.currentWidget()
        if isinstance(widget, QTableWidget):
            return widget
        return None

    # ------------------------------------------------------------------
    def _current_source_name(self) -> str:
        table = self._current_table()
        if table is None:
            return ""
        source = table.property("source")
        return str(source) if source is not None else ""

    # ------------------------------------------------------------------
    def _get_dataframe_for_source(
        self, source: Optional[str] = None, raw: bool = False
    ) -> pd.DataFrame:
        df = self.data_manager.dataframe
        if df.empty:
            return pd.DataFrame()
        key = source or self._current_source_name()
        if key in self._tab_indices:
            indices = self._tab_indices.get(key, [])
            if indices:
                subset = df.loc[indices]
                return subset if raw else subset.reset_index(drop=True).copy()
            return pd.DataFrame()
        return df if raw else df.copy()

    # ------------------------------------------------------------------
    def _default_sheet_name(self, source: str) -> str:
        df = self.data_manager.dataframe
        if df.empty or "SheetName" not in df.columns:
            return "Manual"
        if source in self._tab_indices:
            indices = self._tab_indices.get(source, [])
            if indices:
                subset = df.loc[indices]
                first = str(subset.iloc[0].get("SheetName", "")).strip()
                if first:
                    return first
        if "SourceFile" in df.columns:
            matches = df[df["SourceFile"] == source]
            if not matches.empty:
                first = str(matches.iloc[0].get("SheetName", "")).strip()
                if first:
                    return first
        return "Manual"

    # ------------------------------------------------------------------
    def _create_table_widget(self) -> QTableWidget:
        table = QTableWidget()
        table.setEditTriggers(QTableWidget.EditTrigger.DoubleClicked)
        table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        table.customContextMenuRequested.connect(
            lambda position, table=table: self.table_context_menu(table, position)
        )
        table.itemChanged.connect(self.on_table_cell_changed)
        table.setWordWrap(True)
        header = table.horizontalHeader()
        header.setStretchLastSection(False)
        header.setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        v_header = table.verticalHeader()
        v_header.setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        v_header.setDefaultSectionSize(48)
        return table

    # ------------------------------------------------------------------
    def _calculate_insert_position(self, table: QTableWidget, insert_after: bool) -> int:
        indices = self._tab_indices.get(str(table.property("source") or ""), [])
        current_row = table.currentRow()
        if current_row < 0 or current_row >= len(indices):
            return len(self.data_manager.dataframe)
        base_index = int(indices[current_row])
        return base_index + 1 if insert_after else base_index

    # ------------------------------------------------------------------
    def _insert_blank_row(self, source: str, position: int, source_type: str) -> None:
        df = self.data_manager.dataframe
        if df.empty:
            columns = list(REQUIRED_COLUMNS) + list(OPTIONAL_METADATA_COLUMNS) + [
                self.data_manager.section_column_name
            ]
            df = pd.DataFrame(columns=columns)

        record = {col: "" for col in df.columns}
        record[self.data_manager.section_column_name] = ""
        record["SourceFile"] = source
        if "SheetName" in record:
            record["SheetName"] = self._default_sheet_name(source)
        if "SourceType" in record:
            record["SourceType"] = source_type

        upper = df.iloc[:position]
        lower = df.iloc[position:]
        new_df = pd.concat([upper, pd.DataFrame([record]), lower], ignore_index=True)
        new_df = self.data_manager._apply_section_numbering(new_df)
        self.data_manager.dataframe = new_df
        self._undo_stack.append(new_df.copy())
        self.populate_table()

    # ------------------------------------------------------------------
    def _delete_rows(self, indices: list[int]) -> None:
        if not indices:
            return
        df = self.data_manager.dataframe.drop(index=indices).reset_index(drop=True)
        df = self.data_manager._apply_section_numbering(df)
        self.data_manager.dataframe = df
        self._undo_stack.append(df.copy())
        self.populate_table()

    # ------------------------------------------------------------------
    def remove_current_tab(self) -> None:
        table = self._current_table()
        if table is None:
            QMessageBox.information(self, "Remove Tab", "No tab is currently selected.")
            return
        source = str(table.property("source") or "")
        if not source or source == "No Data":
            QMessageBox.information(self, "Remove Tab", "Cannot remove this tab.")
            return
        indices = list(self._tab_indices.get(source, []))
        if not indices:
            QMessageBox.information(
                self, "Remove Tab", "No rows are associated with this tab."
            )
            return
        reply = QMessageBox.question(
            self,
            "Remove Tab",
            f"Remove all data for tab '{source}'?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No,
        )
        if reply != QMessageBox.StandardButton.Yes:
            return
        self.header_profiles.pop(self._header_key(source), None)
        self._tab_source_types.pop(source, None)
        self._delete_rows(indices)
        self.log_console(f"Removed tab: {source}")

    # ------------------------------------------------------------------
    def _select_global_row(self, source: str, global_index: int) -> None:
        table = self._tab_tables.get(source)
        if table is None:
            return
        indices = self._tab_indices.get(source, [])
        try:
            row = indices.index(global_index)
        except ValueError:
            return
        table.blockSignals(True)
        try:
            table.setCurrentCell(row, 0)
            table.selectRow(row)
            item = table.item(row, 0)
            if item is not None:
                table.scrollToItem(item)
        finally:
            table.blockSignals(False)

    # ------------------------------------------------------------------
    def _select_source_tab(self, source: str) -> None:
        if not source:
            source = "Manual"
        for index in range(self.table_tabs.count()):
            widget = self.table_tabs.widget(index)
            if isinstance(widget, QTableWidget):
                if str(widget.property("source") or "Manual") == source:
                    self.table_tabs.setCurrentIndex(index)
                    break

    # ------------------------------------------------------------------
    def _import_excel_files(self, files: Sequence[str]) -> None:
        normalized = [str(Path(path)) for path in files if path]
        if not normalized:
            return
        try:
            new_df = self.data_manager.load_workbooks(normalized)
        except RequirementDataError as exc:
            LOGGER.exception("Failed to load workbooks")
            QMessageBox.critical(self, "Load Error", str(exc))
            return
        df = self.data_manager.merge_new_dataframe(new_df)
        self._ensure_header_profile("Excel")
        self._undo_stack = [df.copy()]
        self.populate_table()
        self.log_console("Excel files loaded and numbering applied.")

    # ------------------------------------------------------------------
    def _import_word_files(self, files: Sequence[str]) -> None:
        normalized = [str(Path(path)) for path in files if path]
        if not normalized:
            return
        try:
            new_df = self.data_manager.load_word_documents(normalized)
        except RequirementDataError as exc:
            LOGGER.exception("Failed to load Word documents")
            QMessageBox.critical(self, "Load Error", str(exc))
            return
        df = self.data_manager.merge_new_dataframe(new_df)
        if "SourceFile" in new_df.columns:
            for source in new_df["SourceFile"].unique():
                self._ensure_header_profile(str(source))
        self._undo_stack = [df.copy()]
        self.populate_table()
        self.log_console("Word documents loaded and numbering applied.")

    # ------------------------------------------------------------------
    def load_excels(self) -> None:
        files, _ = QFileDialog.getOpenFileNames(
            self, "Select Excel Files", "", "Excel Files (*.xlsx *.xls *.xlsm *.xlsb)"
        )
        if not files:
            return
        self._import_excel_files(files)

    # ------------------------------------------------------------------
    def load_word_documents(self) -> None:
        files, _ = QFileDialog.getOpenFileNames(
            self, "Select Word Files", "", "Word Documents (*.docx *.docm)"
        )
        if not files:
            return
        self._import_word_files(files)

    # ------------------------------------------------------------------
    def dragEnterEvent(self, event) -> None:
        if event.mimeData().hasUrls():
            for url in event.mimeData().urls():
                if not url.isLocalFile():
                    continue
                ext = Path(url.toLocalFile()).suffix.lower()
                if ext in {".xls", ".xlsx", ".xlsm", ".xlsb", ".docx", ".docm"}:
                    event.acceptProposedAction()
                    return
        event.ignore()

    # ------------------------------------------------------------------
    def dropEvent(self, event) -> None:
        if not event.mimeData().hasUrls():
            event.ignore()
            return

        excel_files: list[str] = []
        word_files: list[str] = []
        unsupported: list[str] = []

        for url in event.mimeData().urls():
            if not url.isLocalFile():
                continue
            path = Path(url.toLocalFile())
            if not path.exists():
                continue
            ext = path.suffix.lower()
            if ext in {".xls", ".xlsx", ".xlsm", ".xlsb"}:
                excel_files.append(str(path))
            elif ext in {".docx", ".docm"}:
                word_files.append(str(path))
            elif ext in {".doc"}:
                unsupported.append(str(path))

        if not excel_files and not word_files:
            if unsupported:
                QMessageBox.warning(
                    self,
                    "Unsupported Files",
                    "The following files are not supported:\n" + "\n".join(unsupported),
                )
            event.ignore()
            return

        if excel_files:
            self._import_excel_files(excel_files)
        if word_files:
            self._import_word_files(word_files)
        if unsupported:
            QMessageBox.warning(
                self,
                "Unsupported Files",
                "The following files are not supported:\n" + "\n".join(unsupported),
            )
        event.acceptProposedAction()

    # ------------------------------------------------------------------
    def dragMoveEvent(self, event) -> None:
        self.dragEnterEvent(event)

    # ------------------------------------------------------------------
    def populate_table(self) -> None:
        df = self.data_manager.dataframe
        columns = self.data_manager.visible_columns

        self._loading = True
        self.table_tabs.blockSignals(True)
        self.table_tabs.clear()
        self._tab_tables.clear()
        self._tab_indices.clear()
        self._tab_source_types.clear()

        if df.empty or not columns:
            placeholder = self._create_table_widget()
            placeholder.setEnabled(False)
            placeholder.setRowCount(0)
            placeholder.setColumnCount(len(columns))
            if columns:
                placeholder.setHorizontalHeaderLabels(columns)
            self.table_tabs.addTab(placeholder, "No Data")
            self.table_tabs.blockSignals(False)
            self._loading = False
            self.view_stack.setCurrentWidget(self.table_tabs)
            self.populate_navigation()
            return

        tab_groups: list[tuple[str, pd.DataFrame, str]] = []
        if "SourceType" in df.columns:
            type_series = df["SourceType"].astype(str).str.lower()
            excel_mask = type_series == "excel"
            if excel_mask.any():
                tab_groups.append(("Excel", df[excel_mask], "excel"))
            word_mask = type_series.isin({"word", "doc", "docx", "docm"})
            if word_mask.any():
                for source, group in df[word_mask].groupby("SourceFile", sort=False):
                    tab_groups.append(
                        (str(source).strip() or "Word Document", group, "docx")
                    )
            other_mask = ~(excel_mask | word_mask)
            if other_mask.any():
                for source, group in df[other_mask].groupby("SourceFile", sort=False):
                    stype = (
                        str(group["SourceType"].iloc[0]).lower()
                        if "SourceType" in group.columns
                        else "manual"
                    )
                    tab_groups.append((str(source).strip() or "Data", group, stype))
        else:
            for source, group in df.groupby("SourceFile", sort=False):
                tab_groups.append((str(source).strip() or "Data", group, "manual"))

        if not tab_groups:
            placeholder = self._create_table_widget()
            placeholder.setEnabled(False)
            placeholder.setRowCount(0)
            placeholder.setColumnCount(len(columns))
            if columns:
                placeholder.setHorizontalHeaderLabels(columns)
            self.table_tabs.addTab(placeholder, "No Data")
            self.table_tabs.blockSignals(False)
            self._loading = False
            self.view_stack.setCurrentWidget(self.table_tabs)
            self.populate_navigation()
            return

        for source_name, group, source_type in tab_groups:
            table = self._create_table_widget()
            table.setProperty("source", source_name)
            table.setProperty("source_type", source_type)
            table.setRowCount(len(group))
            table.setColumnCount(len(columns))
            table.setHorizontalHeaderLabels(columns)
            table.horizontalHeader().setStretchLastSection(True)

            row_indices = group.index.tolist()
            self._tab_tables[source_name] = table
            self._tab_indices[source_name] = row_indices
            self._tab_source_types[source_name] = source_type
            self._ensure_header_profile(source_name)

            for display_row, (df_index, row) in enumerate(group.iterrows()):
                for column_index, column_name in enumerate(columns):
                    value = str(row.get(column_name, ""))
                    item = QTableWidgetItem(value)
                    if "heading" in str(row.get("Object Type", "")).lower():
                        font = item.font()
                        font.setBold(True)
                        item.setFont(font)
                    item.setTextAlignment(
                        Qt.AlignmentFlag.AlignTop | Qt.AlignmentFlag.AlignLeft
                    )
                    item.setData(Qt.ItemDataRole.UserRole, df_index)
                    table.setItem(display_row, column_index, item)
            table.resizeColumnsToContents()
            table.resizeRowsToContents()
            self.table_tabs.addTab(table, source_name)

        self.table_tabs.blockSignals(False)
        self._loading = False
        self.view_stack.setCurrentWidget(self.table_tabs)
        self._ensure_header_profile(self._current_source_name())
        self.log_console("Table populated with Excel-style formatting.")
        self.populate_navigation()

    # ------------------------------------------------------------------
    def populate_navigation(self) -> None:
        self.nav_tree.clear()
        subset = self._get_dataframe_for_source(raw=True)
        if subset.empty:
            return
        current_h1: QTreeWidgetItem | None = None
        current_h2: QTreeWidgetItem | None = None
        for global_index, row in subset.iterrows():
            obj_type = str(row.get("Object Type", "")).strip().lower()
            section = str(row.get(self.data_manager.section_column_name, "")).strip()
            text = str(row.get("Object Text", "")).strip()
            if not text:
                continue
            label = f"{section} {text}".strip()
            payload = {
                "index": int(global_index),
                "section": section,
                "text": text,
                "type": obj_type,
                "anchor": f"row-{global_index}",
            }
            if obj_type == "heading 1":
                current_h1 = QTreeWidgetItem([label])
                current_h1.setData(0, Qt.ItemDataRole.UserRole, payload)
                self.nav_tree.addTopLevelItem(current_h1)
                current_h2 = None
            elif obj_type == "heading 2":
                if current_h1 is None:
                    continue
                current_h2 = QTreeWidgetItem([label])
                current_h2.setData(0, Qt.ItemDataRole.UserRole, payload)
                current_h1.addChild(current_h2)
            elif obj_type == "heading 3":
                if current_h2 is None:
                    continue
                item = QTreeWidgetItem([label])
                item.setData(0, Qt.ItemDataRole.UserRole, payload)
                current_h2.addChild(item)
        self.nav_tree.expandAll()

    # ------------------------------------------------------------------
    def on_nav_item_clicked(self, item: QTreeWidgetItem) -> None:
        payload = item.data(0, Qt.ItemDataRole.UserRole) or {}
        global_index = payload.get("index")
        text_value = payload.get("text") or item.text(0).split(" ", 1)[-1].strip()
        section_value = payload.get("section", "")
        anchor_id = payload.get("anchor")
        desired_source = None

        if global_index is not None and global_index in self.data_manager.dataframe.index:
            row = self.data_manager.dataframe.loc[global_index]
            desired_source = str(row.get("SourceFile", "")).strip() or "Manual"
            source_type = str(row.get("SourceType", "")).strip().lower()
            if source_type == "excel" and "Excel" in self._tab_source_types:
                desired_source = "Excel"
            self._select_source_tab(desired_source)
            if self.view_stack.currentWidget() is self.word_preview:
                self.update_word_preview()
            elif self.view_stack.currentWidget() is self.table_tabs:
                self._select_global_row(desired_source, int(global_index))
                table = self._tab_tables.get(desired_source)
                if table:
                    table.setFocus()

        target_candidates = [
            f"{section_value} {text_value}".strip(),
            text_value,
            item.text(0).strip(),
        ]
        if self.view_stack.currentWidget() is self.word_preview:
            def navigate_preview(
                anchor=anchor_id,
                candidates=tuple(target_candidates),
                fallback=text_value,
            ) -> None:
                if anchor:
                    self.word_preview.scrollToAnchor(anchor)
                matched = self._highlight_preview_text(candidates)
                if matched:
                    self.log_console(f"Navigated to: {matched}")
                else:
                    self.log_console(f"Could not find heading: {fallback}")

            QTimer.singleShot(0, navigate_preview)
            return

        self.log_console(f"Navigated to: {text_value}")

    # ------------------------------------------------------------------
    def on_table_tab_changed(self, index: int) -> None:
        if self._loading:
            return
        self._ensure_header_profile(self._current_source_name())
        self.populate_navigation()
        if (
            self.view_stack.currentWidget() is self.word_preview
            and not self._get_dataframe_for_source().empty
        ):
            self.update_word_preview()

    # ------------------------------------------------------------------
    def on_table_cell_changed(self, item: QTableWidgetItem) -> None:
        if self._loading:
            return
        table = item.tableWidget()
        if not isinstance(table, QTableWidget):
            return
        original_index = item.data(Qt.ItemDataRole.UserRole)
        if original_index is None:
            return
        column_header = table.horizontalHeaderItem(item.column())
        if column_header is None:
            return
        column_name = column_header.text()

        try:
            self.data_manager.update_cell(int(original_index), column_name, item.text().strip())
        except RequirementDataError as exc:
            QMessageBox.warning(self, "Edit Error", str(exc))
            self._loading = True
            try:
                previous_value = self.data_manager.dataframe.at[int(original_index), column_name]
                item.setText(str(previous_value))
            finally:
                self._loading = False
            return

        self.data_manager.dataframe = self.data_manager._apply_section_numbering(
            self.data_manager.dataframe
        )
        self._undo_stack.append(self.data_manager.dataframe.copy())
        self.populate_table()

    # ------------------------------------------------------------------
    def table_context_menu(self, table: QTableWidget, position) -> None:
        if table is None:
            return
        menu = QMenu(self)
        add_row = menu.addAction("Add Row Below")
        delete_row = menu.addAction("Delete Row")
        action = menu.exec(table.viewport().mapToGlobal(position))

        if action is None:
            return

        source = str(table.property("source") or "")
        source_type = str(
            table.property("source_type")
            or self._tab_source_types.get(source, "manual")
        )
        if action == add_row:
            insert_at = self._calculate_insert_position(table, insert_after=True)
            self._insert_blank_row(source, insert_at, source_type)
        elif action == delete_row:
            indices = self._tab_indices.get(source, [])
            current_row = table.currentRow()
            if 0 <= current_row < len(indices):
                self._delete_rows([indices[current_row]])

    # ------------------------------------------------------------------
    def add_image_attachment(self) -> None:
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Image",
            "",
            "Image Files (*.png *.jpg *.jpeg *.bmp *.gif *.tif *.tiff)",
        )
        if not file_path:
            return

        try:
            data = Path(file_path).read_bytes()
        except OSError as exc:
            QMessageBox.critical(self, "Image Error", f"Could not read image: {exc}")
            return

        mime = mimetypes.guess_type(file_path)[0] or "image/png"
        payload = {
            "mime": mime,
            "data": base64.b64encode(data).decode("ascii"),
            "filename": Path(file_path).name,
        }

        table = self._current_table()
        source = self._current_source_name() or "Manual"
        source_type = str(
            (table.property("source_type") if table else None)
            or self._tab_source_types.get(source, "manual")
        )
        insert_position = (
            self._calculate_insert_position(table, insert_after=True)
            if table is not None
            else len(self.data_manager.dataframe)
        )

        df = self.data_manager.insert_attachment(
            object_type="Image",
            attachment_type="image",
            attachment_data=json.dumps(payload),
            object_text=Path(file_path).name,
            insert_at=insert_position,
            source_file=source,
            sheet_name=self._default_sheet_name(source),
            source_type=source_type,
        )
        self._undo_stack.append(df.copy())
        self.populate_table()
        self._ensure_header_profile(source)
        self._select_global_row(source, insert_position)
        self.log_console(f"Image attachment added: {file_path}")

    # ------------------------------------------------------------------
    def add_table_attachment(self) -> None:
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Table File",
            "",
            "Table Files (*.xlsx *.xls *.csv)",
        )
        if not file_path:
            return

        suffix = Path(file_path).suffix.lower()
        try:
            if suffix in {".xlsx", ".xls"}:
                table_df = pd.read_excel(file_path)
            else:
                table_df = pd.read_csv(file_path)
        except Exception as exc:
            QMessageBox.critical(self, "Table Error", f"Could not load table: {exc}")
            return

        if table_df.empty:
            QMessageBox.information(self, "Empty Table", "Selected table has no data.")
            return

        caption, accepted = QInputDialog.getText(
            self, "Table Caption", "Caption (optional):"
        )
        if not accepted:
            return

        html_table = self.data_manager.dataframe_to_html_table(table_df)

        table = self._current_table()
        source = self._current_source_name() or "Manual"
        source_type = str(
            (table.property("source_type") if table else None)
            or self._tab_source_types.get(source, "manual")
        )
        insert_position = (
            self._calculate_insert_position(table, insert_after=True)
            if table is not None
            else len(self.data_manager.dataframe)
        )

        df = self.data_manager.insert_attachment(
            object_type="Table",
            attachment_type="table",
            attachment_data=html_table,
            object_text=caption.strip(),
            insert_at=insert_position,
            source_file=source,
            sheet_name=self._default_sheet_name(source),
            source_type=source_type,
        )
        self._undo_stack.append(df.copy())
        self.populate_table()
        self._ensure_header_profile(source)
        self._select_global_row(source, insert_position)
        self.log_console(f"Table attachment added: {file_path}")

    # ------------------------------------------------------------------
    def edit_header_details(self) -> None:
        source = self._current_source_name()
        self._ensure_header_profile(source)
        current_settings = self._get_header_settings(source)
        dialog = HeaderSettingsDialog(self, copy.deepcopy(current_settings))
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return
        current_settings.update(dialog.values())
        self.log_console("Header details updated.")
        if self.view_stack.currentWidget() is self.word_preview:
            df = self._get_dataframe_for_source(source, raw=True)
            self.word_preview.setHtml(self.compose_preview_html(df))

    # ------------------------------------------------------------------
    def compose_preview_html(self, df: Optional[pd.DataFrame] = None) -> str:
        df_raw = df.copy() if df is not None else self._get_dataframe_for_source(raw=True)
        if df_raw.empty:
            local_df = pd.DataFrame()
        else:
            local_df = df_raw.reset_index().rename(columns={"index": "_global_index"})
        settings = self._get_header_settings()

        styles = """
        <style>
            body { font-family: Arial, sans-serif; font-size: 11pt; color: #000; margin: 20px; }
            .front-page { border: 2px solid #000; padding: 28px; margin-bottom: 36px; }
            .front-title { text-align: center; font-size: 18pt; font-weight: bold; margin-bottom: 12px; }
            .meta-lines { text-align: center; font-size: 12pt; margin-bottom: 16px; }
            .signature-table { width: 100%; border-collapse: collapse; margin-bottom: 18px; }
            .signature-table td, .signature-table th { border: 1px solid #000; padding: 6px; vertical-align: top; }
            .signature-table th { text-align: left; width: 22%; font-weight: bold; }
            .signature-name { font-weight: bold; }
            .signature-title { font-style: italic; font-size: 10pt; }
            .logo-row { display: flex; justify-content: space-between; align-items: center; margin: 20px 0; }
            .logo-cell { width: 48%; text-align: center; }
            .logo-cell img { max-height: 90px; }
            .address-row { display: flex; justify-content: space-between; }
            .address-block { width: 48%; text-align: left; line-height: 1.4; }
            .notice { border: 1px solid #000; padding: 12px; margin-top: 24px; }
            .notice-title { font-weight: bold; text-align: center; margin-bottom: 8px; }
            .copyright { text-align: center; margin-top: 18px; font-size: 10pt; }
            .page-separator { border-top: 2px solid #555; margin: 36px 0; }
            .toc-section { margin-bottom: 24px; }
            .toc-section h2 { font-size: 14pt; margin-bottom: 8px; }
            .toc-item { margin-left: 8px; line-height: 1.4; }
            .toc-item.level-2 { margin-left: 24px; }
            .toc-item.level-3 { margin-left: 40px; }
            .toc-item a { color: #000; text-decoration: none; }
            .toc-item a:hover { text-decoration: underline; }
            .requirement-id { font-weight: bold; margin-top: 16px; margin-bottom: 4px; }
            .requirement-text { margin-top: 0; text-align: left; line-height: 1.5; }
            .body-text { text-align: left; line-height: 1.5; margin-bottom: 12px; }
            table { margin: 12px auto; border-collapse: collapse; width: 100%; }
            table, td, th { border: 1px solid #000; padding: 4px; }
            .table-block { text-align: center; margin: 16px 0; }
            .image-block { text-align: center; margin: 16px 0; }
            .caption { font-style: italic; font-size: 10pt; margin-top: 6px; text-align: center; }
        </style>
        """

        front_page_html = self._build_front_page_html(settings)

        if local_df.empty:
            return f"<html><head>{styles}</head><body>{front_page_html}<p>No data loaded.</p></body></html>"

        toc_entries = []
        figure_entries = []
        table_entries = []
        body_parts: list[str] = []
        figure_count = 0
        table_count = 0
        section_col = self.data_manager.section_column_name

        for idx, row in local_df.iterrows():
            global_idx = row.get("_global_index", idx)
            obj_type = str(row.get("Object Type", "")).strip().lower()
            section = str(row.get(section_col, "")).strip()
            text = str(row.get("Object Text", "")).strip()
            req_id = str(row.get("Requirement ID", "")).strip()
            attachment_type = str(row.get("Attachment Type", "")).strip().lower()
            attachment_data = str(row.get("Attachment Data", "")).strip()

            if obj_type.startswith("heading"):
                level = int(obj_type.split()[-1])
                heading_text = f"{section} {text}".strip() or text or section or "Heading"
                heading_id = f"row-{global_idx}"
                body_parts.append(
                    f'<h{level} id="{heading_id}">{html.escape(heading_text)}</h{level}>'
                )
                toc_entries.append(
                    {"level": level, "id": heading_id, "title": heading_text}
                )
                continue

            if attachment_type == "table" and attachment_data:
                table_count += 1
                table_id = f"table-{global_idx}"
                caption_text = text or f"Table {table_count}"
                body_parts.append(f'<div id="{table_id}" class="table-block">{attachment_data}')
                body_parts.append(
                    f'<div class="caption">Table {table_count}: {html.escape(caption_text)}</div></div>'
                )
                table_entries.append(
                    {"id": table_id, "title": f"Table {table_count}: {caption_text}"}
                )
                continue

            if attachment_type == "image" and attachment_data:
                figure_count += 1
                figure_id = f"figure-{global_idx}"
                try:
                    payload = json.loads(attachment_data)
                except json.JSONDecodeError:
                    payload = {"data": attachment_data, "mime": "image/png", "filename": text or "Image"}
                image_data = payload.get("data", "")
                if not image_data:
                    continue
                mime = payload.get("mime", "image/png")
                caption_text = text or payload.get("filename", "") or f"Figure {figure_count}"
                width_percent = settings.get("preview_image_width_percent", 80)
                body_parts.append(
                    f'<div id="{figure_id}" class="image-block"><img src="data:{mime};base64,{image_data}" '
                    f'alt="{html.escape(caption_text)}" style="max-width:{width_percent}%; width:{width_percent}%; height:auto;"/></div>'
                )
                body_parts.append(
                    f'<div class="caption">Figure {figure_count}: {html.escape(caption_text)}</div>'
                )
                figure_entries.append(
                    {"id": figure_id, "title": f"Figure {figure_count}: {caption_text}"}
                )
                continue

            anchor_id = f"row-{global_idx}"
            if req_id:
                requirement_text = html.escape(text) if text else "&nbsp;"
                body_parts.append(
                    f'<div id="{anchor_id}" class="requirement-block">'
                    f'<p class="requirement-id">Requirement ID: {html.escape(req_id)}</p>'
                    f'<p class="requirement-text">{requirement_text}</p>'
                    "</div>"
                )
                continue

            if text:
                body_parts.append(
                    f'<p id="{anchor_id}" class="body-text">{html.escape(text)}</p>'
                )

        toc_html = ""
        if toc_entries:
            toc_html = '<div class="toc-section"><h2>Table of Contents</h2>'
            for entry in toc_entries:
                toc_html += (
                    f'<div class="toc-item level-{entry["level"]}">'
                    f'<a href="#{entry["id"]}">{html.escape(entry["title"])}</a>'
                    "</div>"
                )
            toc_html += "</div>"

        figures_html = ""
        if figure_entries:
            figures_html = '<div class="toc-section"><h2>List of Figures</h2><ol>'
            for entry in figure_entries:
                figures_html += (
                    f'<li><a href="#{entry["id"]}">{html.escape(entry["title"])}</a></li>'
                )
            figures_html += "</ol></div>"

        tables_html = ""
        if table_entries:
            tables_html = '<div class="toc-section"><h2>List of Tables</h2><ol>'
            for entry in table_entries:
                tables_html += (
                    f'<li><a href="#{entry["id"]}">{html.escape(entry["title"])}</a></li>'
                )
            tables_html += "</ol></div>"

        body_html = "\n".join(body_parts) if body_parts else "<p>No requirement content available.</p>"

        return (
            f"<html><head>{styles}</head><body>"
            f"{front_page_html}"
            '<div class="page-separator"></div>'
            f"{toc_html}{figures_html}{tables_html}{body_html}"
            "</body></html>"
        )

    # ------------------------------------------------------------------
    def _highlight_preview_text(self, candidates: Sequence[str]) -> Optional[str]:
        cursor = self.word_preview.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.End)
        self.word_preview.setTextCursor(cursor)
        for candidate in dict.fromkeys([c for c in candidates if c]):
            if self.word_preview.find(
                candidate, QTextDocument.FindFlag.FindBackward
            ):
                self.word_preview.ensureCursorVisible()
                return candidate
        return None

    # ------------------------------------------------------------------
    def _build_front_page_html(self, settings: Dict[str, object]) -> str:
        title = html.escape(str(settings.get("document_title", "Requirement Document")))
        doc_number = html.escape(str(settings.get("document_number", "")))
        revision = html.escape(str(settings.get("revision", "")))

        meta_lines = []
        if doc_number:
            meta_lines.append(f"Document No: {doc_number}")
        if revision:
            meta_lines.append(f"Revision: {revision}")
        meta_html = "<br/>".join(meta_lines)

        rows = []
        people = [
            ("Author", "author_name", "author_title"),
            ("Reviewer", "reviewer_name", "reviewer_title"),
            ("Quality Assurance", "qa_name", "qa_title"),
            ("Configuration Manager", "config_manager_name", "config_manager_title"),
        ]
        for label, name_key, title_key in people:
            name = str(settings.get(name_key, "")).strip()
            title_text = str(settings.get(title_key, "")).strip()
            name_html = html.escape(name) if name else "&nbsp;"
            title_html = html.escape(title_text) if title_text else ""
            cell_content = (
                f'<div class="signature-name">{name_html}</div>'
                + (f'<div class="signature-title">{title_html}</div>' if title_text else "")
            )
            rows.append(
                f"<tr><th>{label}:</th><td>{cell_content}</td><td>&nbsp;</td></tr>"
            )

        signature_table = (
            '<table class="signature-table">'
            + "".join(rows)
            + "</table>"
        )

        left_logo = self._logo_img_tag(str(settings.get("logo_left_path", "")), "Left Logo")
        right_logo = self._logo_img_tag(str(settings.get("logo_right_path", "")), "Right Logo")
        logo_row = ""
        if left_logo or right_logo:
            logo_row = (
                '<div class="logo-row">'
                f'<div class="logo-cell">{left_logo or ""}</div>'
                f'<div class="logo-cell">{right_logo or ""}</div>'
                "</div>"
            )

        address_howell = str(settings.get("address_howell", ""))
        address_mach = str(settings.get("address_mach", ""))
        addresses_html = ""
        if address_howell or address_mach:
            addresses_html = (
                '<div class="address-row">'
                f'<div class="address-block">{address_howell}</div>'
                f'<div class="address-block">{address_mach}</div>'
                "</div>"
            )

        notice_text = html.escape(str(settings.get("proprietary_notice", ""))).replace(
            "\n", "<br/>"
        )
        notice_html = (
            '<div class="notice">'
            '<div class="notice-title">Proprietary Notice</div>'
            f'<div class="notice-body">{notice_text}</div>'
            "</div>"
        )

        copyright_text = html.escape(str(settings.get("copyright_notice", "")))

        front_parts = [
            '<div class="front-page">',
            f'<div class="front-title">{title}</div>',
        ]
        if meta_html:
            front_parts.append(f'<div class="meta-lines">{meta_html}</div>')
        front_parts.append(signature_table)
        if logo_row:
            front_parts.append(logo_row)
        if addresses_html:
            front_parts.append(addresses_html)
        if copyright_text:
            front_parts.append(f'<div class="copyright">{copyright_text}</div>')
        front_parts.append(notice_html)
        front_parts.append("</div>")
        return "".join(front_parts)

    # ------------------------------------------------------------------
    def _logo_img_tag(self, path: str, alt: str) -> str:
        if not path:
            return ""
        try:
            data = Path(path).read_bytes()
        except OSError:
            return ""
        mime = mimetypes.guess_type(path)[0] or "image/png"
        encoded = base64.b64encode(data).decode("ascii")
        return (
            f'<img src="data:{mime};base64,{encoded}" alt="{html.escape(alt)}" />'
        )

    # ------------------------------------------------------------------
    def update_word_preview(self) -> None:
        html_content = self.compose_preview_html()
        self.word_preview.setHtml(html_content)

    # ------------------------------------------------------------------
    def undo_last(self) -> None:
        if len(self._undo_stack) <= 1:
            return
        self._undo_stack.pop()
        self.data_manager.dataframe = self._undo_stack[-1].copy()
        self.populate_table()
        self.populate_navigation()
        self.log_console("Undo applied.")

    # ------------------------------------------------------------------
    def run_convert_preview(self) -> None:
        df = self._get_dataframe_for_source(raw=True)
        if df.empty:
            QMessageBox.information(self, "Empty", "No data available for preview.")
            return
        html_content = self.compose_preview_html(df)
        self.word_preview.setHtml(html_content)
        self.view_stack.setCurrentIndex(1)
        self.log_console("Word preview generated.")

    # ------------------------------------------------------------------
    def show_table_view(self) -> None:
        self.view_stack.setCurrentWidget(self.table_tabs)
        self.log_console("Returned to table view.")

    # ------------------------------------------------------------------
    def save_word(self) -> None:
        df = self._get_dataframe_for_source(raw=True)
        if df.empty:
            QMessageBox.warning(self, "Empty", "Nothing to save!")
            return

        file_name, _ = QFileDialog.getSaveFileName(
            self, "Save Word File", "", "Word Document (*.docx)"
        )
        if not file_name:
            return
        if not file_name.endswith(".docx"):
            file_name += ".docx"

        try:
            from docx import Document
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.text import (
                WD_ALIGN_PARAGRAPH,
                WD_TAB_ALIGNMENT,
                WD_TAB_LEADER,
            )
            from docx.oxml import OxmlElement, parse_xml
            from docx.oxml.ns import qn, nsdecls
            from docx.shared import Inches, Pt, RGBColor
            from bs4 import BeautifulSoup
        except Exception as exc:  # pragma: no cover - optional dependency
            QMessageBox.critical(
                self,
                "Dependency Error",
                "python-docx and beautifulsoup4 are required to export Word files.\n"
                f"Details: {exc}",
            )
            return

        settings = self._get_header_settings()
        document = Document()
        self._configure_document_template(
            document, settings, Pt, Inches, OxmlElement, qn
        )
        self._build_header_footer(
            document,
            settings,
            WD_ALIGN_PARAGRAPH,
            WD_TAB_ALIGNMENT,
            WD_TAB_LEADER,
            Pt,
            Inches,
            OxmlElement,
            qn,
            parse_xml,
            nsdecls,
        )
        self._add_front_matter(
            document,
            settings,
            WD_ALIGN_PARAGRAPH,
            Pt,
            Inches,
            WD_TABLE_ALIGNMENT,
            OxmlElement,
            qn,
        )
        document.add_page_break()
        self._add_table_of_contents(document, settings, OxmlElement, qn)
        document.add_page_break()
        self._append_body_content(
            document,
            df,
            settings,
            BeautifulSoup,
            WD_ALIGN_PARAGRAPH,
            Inches,
            WD_TABLE_ALIGNMENT.CENTER,
            OxmlElement,
            qn,
            RGBColor,
        )
        document.save(file_name)
        self.log_console(
            f"Word file saved: {file_name} (update fields in Word to refresh TOC)"
        )

    # ------------------------------------------------------------------
    def _configure_document_template(
        self, document, settings, Pt, Inches, OxmlElement, qn
    ) -> None:
        section = document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        self._apply_page_border(section, OxmlElement, qn)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(11)

        for style_name, size in (("Heading 1", 12), ("Heading 2", 12), ("Heading 3", 11)):
            style = document.styles[style_name]
            style.font.name = "Arial"
            style.font.size = Pt(size)
            style.font.bold = True

    # ------------------------------------------------------------------
    def _apply_page_border(self, section, OxmlElement, qn) -> None:
        sect_pr = section._sectPr
        existing = sect_pr.find(qn("w:pgBorders"))
        if existing is not None:
            sect_pr.remove(existing)

        page_borders = OxmlElement("w:pgBorders")
        page_borders.set(qn("w:offsetFrom"), "page")
        for edge in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "12")
            border.set(qn("w:space"), "24")
            border.set(qn("w:color"), "000000")
            page_borders.append(border)
        sect_pr.append(page_borders)

    # ------------------------------------------------------------------
    def _build_header_footer(
        self,
        document,
        settings,
        WD_ALIGN_PARAGRAPH,
        WD_TAB_ALIGNMENT,
        WD_TAB_LEADER,
        Pt,
        Inches,
        OxmlElement,
        qn,
        parse_xml,
        nsdecls,
    ) -> None:
        section = document.sections[0]

        header = section.header
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0]
        header_para.text = ""

        doc_number = str(settings.get("document_number", ""))
        title = str(settings.get("document_title", ""))
        revision = str(settings.get("revision", ""))
        core_parts = [part for part in (doc_number, title) if part]
        if revision:
            header_text = " \u2013 ".join(core_parts)
            if header_text:
                header_text = f"{header_text} - Revision: {revision}"
            else:
                header_text = f"Revision: {revision}"
        else:
            header_text = " \u2013 ".join(core_parts)
        if not header_text:
            header_text = "Requirement Document"
        header_run = header_para.add_run(header_text)
        header_run.font.name = "Arial"
        header_run.font.size = Pt(10)
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0]
        footer_para.text = ""
        paragraph_format = footer_para.paragraph_format
        paragraph_format.tab_stops.clear_all()
        paragraph_format.tab_stops.add_tab_stop(
            Inches(3.5), WD_TAB_ALIGNMENT.CENTER, WD_TAB_LEADER.SPACES
        )
        paragraph_format.tab_stops.add_tab_stop(
            Inches(7.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES
        )
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(0)

        left_run = footer_para.add_run("Mach Global Technologies       ")
        left_run.font.name = "Arial"
        footer_para.add_run("\t")
        center_run = footer_para.add_run("Howell Instruments Proprietary Information")
        center_run.font.name = "Arial"
        footer_para.add_run("\tPage ")
        self._append_page_field(footer_para, OxmlElement, qn)
        watermark_text = str(settings.get("watermark_text", "")).strip()
        if watermark_text:
            self._apply_watermark(document, watermark_text, parse_xml, nsdecls)


    # ------------------------------------------------------------------
    def _append_page_field(self, paragraph, OxmlElement, qn) -> None:
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        paragraph._p.append(field)

    # ------------------------------------------------------------------
    def _add_front_matter(
        self,
        document,
        settings,
        WD_ALIGN_PARAGRAPH,
        Pt,
        Inches,
        WD_TABLE_ALIGNMENT,
        OxmlElement,
        qn,
    ) -> None:
        title = settings.get("document_title") or "Requirement Document"
        doc_number = settings.get("document_number")
        revision = settings.get("revision")

        title_para = document.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.name = "Arial"
        title_run.font.size = Pt(18)

        if doc_number:
            doc_para = document.add_paragraph()
            doc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc_run = doc_para.add_run(f"Document No: {doc_number}")
            doc_run.font.name = "Arial"
            doc_run.font.size = Pt(12)

        if revision:
            rev_para = document.add_paragraph()
            rev_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rev_run = rev_para.add_run(f"Revision: {revision}")
            rev_run.font.name = "Arial"
            rev_run.font.size = Pt(12)

        document.add_paragraph()

        roles = [
            ("Author", "author_name", "author_title"),
            ("Reviewer", "reviewer_name", "reviewer_title"),
            ("Quality Assurance", "qa_name", "qa_title"),
            ("Configuration Manager", "config_manager_name", "config_manager_title"),
        ]
        table = document.add_table(rows=len(roles), cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_idx, (label, name_key, title_key) in enumerate(roles):
            cells = table.rows[row_idx].cells
            cells[0].text = f"{label}:"
            for run in cells[0].paragraphs[0].runs:
                run.font.bold = True
            name = str(settings.get(name_key, "")).strip()
            title_text = str(settings.get(title_key, "")).strip()
            cell_lines = []
            if name:
                cell_lines.append(name)
            if title_text:
                cell_lines.append(title_text)
            cells[1].text = "\n".join(cell_lines)
            cells[2].text = ""

        document.add_paragraph()

        logo_table = document.add_table(rows=1, cols=2)
        logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._clear_table_borders(logo_table, OxmlElement, qn)
        for idx, key in enumerate(("logo_left_path", "logo_right_path")):
            cell = logo_table.cell(0, idx)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._insert_logo(
                paragraph,
                str(settings.get(key, "")),
                Inches,
            )

        document.add_paragraph()

        address_table = document.add_table(rows=1, cols=2)
        address_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._clear_table_borders(address_table, OxmlElement, qn)
        address_howell = str(settings.get("address_howell", "")).replace("<br/>", "\n")
        address_mach = str(settings.get("address_mach", "")).replace("<br/>", "\n")
        address_table.cell(0, 0).text = address_howell
        address_table.cell(0, 1).text = address_mach

        document.add_paragraph()

        copyright_text = str(settings.get("copyright_notice", "")).strip()
        if copyright_text:
            copyright_para = document.add_paragraph(copyright_text)
            copyright_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        notice_table = document.add_table(rows=2, cols=1)
        notice_table.style = "Table Grid"
        notice_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        notice_table.cell(0, 0).text = "Proprietary Notice"
        notice_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in notice_table.cell(0, 0).paragraphs[0].runs:
            run.font.bold = True
        notice_text = str(settings.get("proprietary_notice", "")).replace("<br/>", "\n")
        notice_table.cell(1, 0).text = notice_text

        document.add_paragraph()

    # ------------------------------------------------------------------
    def _add_table_of_contents(self, document, settings, OxmlElement, qn) -> None:
        document.add_paragraph("Table of Contents", style="Heading 1")
        toc_para = document.add_paragraph()
        toc_field = OxmlElement("w:fldSimple")
        toc_field.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
        toc_para._p.append(toc_field)

        document.add_paragraph()
        document.add_paragraph("List of Figures", style="Heading 1")
        lof_para = document.add_paragraph()
        lof_field = OxmlElement("w:fldSimple")
        lof_field.set(qn("w:instr"), 'TOC \\h \\z \\c "Figure"')
        lof_para._p.append(lof_field)

        document.add_paragraph()
        document.add_paragraph("List of Tables", style="Heading 1")
        lot_para = document.add_paragraph()
        lot_field = OxmlElement("w:fldSimple")
        lot_field.set(qn("w:instr"), 'TOC \\h \\z \\c "Table"')
        lot_para._p.append(lot_field)

    # ------------------------------------------------------------------
    def _append_body_content(
        self,
        document,
        df: pd.DataFrame,
        settings,
        BeautifulSoup,
        WD_ALIGN_PARAGRAPH,
        Inches,
        table_alignment,
        OxmlElement,
        qn,
        RGBColor,
    ) -> None:
        section_col = self.data_manager.section_column_name
        for index, row in df.iterrows():
            global_index = int(row.get("_global_index", index)) if "_global_index" in row else index
            obj_type = str(row.get("Object Type", "")).strip().lower()
            section = str(row.get(section_col, "")).strip()
            text_value = str(row.get("Object Text", "")).strip()
            req_id = str(row.get("Requirement ID", "")).strip()
            attachment_type = str(row.get("Attachment Type", "")).strip().lower()
            attachment_data = str(row.get("Attachment Data", "")).strip()

            if obj_type.startswith("heading"):
                try:
                    level = max(1, min(int(obj_type.split()[-1]), 9))
                except ValueError:
                    level = 1
                heading_text = f"{section} {text_value}".strip() or text_value or section or "Heading"
                document.add_heading(heading_text, level=level)
                self._add_horizontal_rule(document, OxmlElement, qn)
                continue

            if attachment_type == "table" and attachment_data:
                self._add_table_from_html(
                    document,
                    attachment_data,
                    text_value,
                    BeautifulSoup,
                    WD_ALIGN_PARAGRAPH,
                    table_alignment,
                    RGBColor,
                    OxmlElement,
                    qn,
                )
                self._add_horizontal_rule(document, OxmlElement, qn)
                continue

            if attachment_type == "image" and attachment_data:
                self._add_image_from_payload(
                    document,
                    attachment_data,
                    text_value,
                    settings,
                    Inches,
                    WD_ALIGN_PARAGRAPH,
                    RGBColor,
                    OxmlElement,
                    qn,
                )
                self._add_horizontal_rule(document, OxmlElement, qn)
                continue

            if req_id:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                label_run = paragraph.add_run("Requirement ID: ")
                label_run.bold = True
                value_run = paragraph.add_run(req_id)
                value_run.bold = False
                if text_value:
                    text_para = document.add_paragraph(text_value)
                    text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self._add_horizontal_rule(document, OxmlElement, qn)
                continue

            if text_value:
                paragraph = document.add_paragraph(text_value)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self._add_horizontal_rule(document, OxmlElement, qn)

    # ------------------------------------------------------------------


    def _add_table_from_html(
        self,
        document,
        html_table: str,
        caption: str,
        BeautifulSoup,
        WD_ALIGN_PARAGRAPH,
        table_alignment,
        RGBColor,
        OxmlElement,
        qn,
    ) -> None:
        soup = BeautifulSoup(html_table, "html.parser")
        rows = soup.find_all("tr")
        if not rows:
            return
        max_cols = max(len(row.find_all(["td", "th"])) for row in rows)
        table = document.add_table(rows=len(rows), cols=max_cols)
        table.style = "Table Grid"
        if table_alignment is not None:
            table.alignment = table_alignment

        for row_idx, row_tag in enumerate(rows):
            cells = row_tag.find_all(["td", "th"])
            for col_idx, cell in enumerate(cells):
                if col_idx >= max_cols:
                    break
                cell_text = cell.get_text(separator="\n").strip()
                table.cell(row_idx, col_idx).text = cell_text

        caption_text = caption.strip()
        caption_para = document.add_paragraph(style="Caption")
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = caption_para.add_run("Table ")
        label_run.font.bold = False
        label_run.font.color.rgb = RGBColor(0, 0, 0)
        self._append_seq_field(caption_para, "Table", RGBColor, OxmlElement, qn)
        tail_text = f": {caption_text}" if caption_text else ""
        tail_run = caption_para.add_run(tail_text)
        tail_run.font.bold = False
        tail_run.font.color.rgb = RGBColor(0, 0, 0)

    def _add_image_from_payload(
        self,
        document,
        payload: str,
        caption: str,
        settings,
        Inches,
        WD_ALIGN_PARAGRAPH,
        RGBColor,
        OxmlElement,
        qn,
    ) -> None:
        try:
            data = json.loads(payload)
        except json.JSONDecodeError:
            data = {"data": payload, "mime": "image/png", "filename": "image"}

        raw_data = data.get("data", "")
        if not raw_data:
            return
        try:
            image_bytes = base64.b64decode(raw_data)
        except (ValueError, TypeError):
            return

        image_stream = io.BytesIO(image_bytes)
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        try:
            width = float(settings.get("export_image_width_inches", 5.5))
            run.add_picture(image_stream, width=Inches(width))
        except Exception:
            return

        caption_text = caption.strip() if caption else data.get("filename", "")
        caption_para = document.add_paragraph(style="Caption")
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = caption_para.add_run("Figure ")
        label_run.font.bold = False
        label_run.font.color.rgb = RGBColor(0, 0, 0)
        self._append_seq_field(caption_para, "Figure", RGBColor, OxmlElement, qn)
        tail_text = f": {caption_text or 'Figure'}"
        tail_run = caption_para.add_run(tail_text)
        tail_run.font.bold = False
        tail_run.font.color.rgb = RGBColor(0, 0, 0)


def _apply_watermark(self, document, text: str, parse_xml, nsdecls) -> None:
    safe_text = (text or "").strip()
    if not safe_text:
        return
    safe_text = html.escape(safe_text).replace('"', "'")
    section = document.sections[0]
    header = section.header
    nsmap = {"v": "urn:schemas-microsoft-com:vml"}
    for shape in header._element.xpath(".//v:shape[contains(@id,'PowerPlusWaterMarkObject')]", namespaces=nsmap):
        parent = shape.getparent()
        if parent is not None:
            parent.remove(shape)
    for shapetype in header._element.xpath(".//v:shapetype[@id='_x0000_t136']", namespaces=nsmap):
        parent = shapetype.getparent()
        if parent is not None:
            parent.remove(shapetype)

    watermark_xml = (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:x="urn:schemas-microsoft-com:office:excel" xmlns:w10="urn:schemas-microsoft-com:office:word">'
        '<w:r><w:pict>'
        '<v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" o:connecttype="custom" '
        'path="m@7,l@8,m@5,21600l@6,21600e" filled="f" stroked="f">'
        '<v:formulas>'
        '<v:f eqn="if lineDrawn pixelLineWidth 0"/>'
        '<v:f eqn="sum @0 1 0"/>'
        '<v:f eqn="sum 0 0 @1"/>'
        '<v:f eqn="prod @2 1 2"/>'
        '<v:f eqn="prod @3 21600 pixelWidth"/>'
        '<v:f eqn="prod @3 21600 pixelHeight"/>'
        '<v:f eqn="sum @0 0 1"/>'
        '<v:f eqn="prod @6 1 2"/>'
        '<v:f eqn="prod @7 21600 pixelWidth"/>'
        '<v:f eqn="prod @7 21600 pixelHeight"/>'
        '</v:formulas>'
        '<v:path o:extrusionok="f" gradientshapeok="t" o:connecttype="custom" '
        'connectlocs="@5,0;@6,21600;@8,0;@9,21600"/>'
        '<o:lock v:ext="edit" text="t" shapetype="t"/>'
        '</v:shapetype>'
        '<v:shape id="PowerPlusWaterMarkObject" o:spid="_x0000_s1025" type="#_x0000_t136" '
        'style="position:absolute;margin-left:0;margin-top:0;width:468pt;height:468pt;rotation:315;z-index:-251658240;visibility:visible;mso-wrap-edited:f" '
        'o:allowincell="f" fillcolor="gray" stroked="f">'
        '<v:fill opacity="0.1" />'
        f"<v:textpath style=\"font-family:'Calibri';font-size:48pt\" string=\"{safe_text}\"/>"
        '</v:shape>'
        '</w:pict></w:r></w:p>'
    )
    header._element.append(parse_xml(watermark_xml))

    def _append_seq_field(self, paragraph, label: str, RGBColor, OxmlElement, qn) -> None:
        run = paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_char_begin)

        instr_text = OxmlElement('w:instrText')
        instr_text.text = f'SEQ {label} \\* ARABIC'
        run._r.append(instr_text)

        fld_char_separate = OxmlElement('w:fldChar')
        fld_char_separate.set(qn('w:fldCharType'), 'separate')
        run._r.append(fld_char_separate)

        value_run = paragraph.add_run('1')
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        value_run._r.append(fld_char_end)

        for seq_run in (run, value_run):
            seq_run.font.bold = False
            seq_run.font.color.rgb = RGBColor(0, 0, 0)

    def _clear_table_borders(self, table, OxmlElement, qn) -> None:
        tbl_pr = table._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            table._tbl.tblPr = tbl_pr
        borders = tbl_pr.find(qn("w:tblBorders"))
        if borders is not None:
            tbl_pr.remove(borders)

    # ------------------------------------------------------------------
    def _insert_logo(self, paragraph, path: str, Inches) -> None:
        if not path:
            return
        logo_path = Path(path)
        if not logo_path.exists():
            return
        try:
            run = paragraph.add_run()
            run.add_picture(str(logo_path), width=Inches(2.3))
        except Exception:
            paragraph.add_run("")

    # ------------------------------------------------------------------
    def navigate_to_requirement(
        self, requirement_id: str, source: Optional[str] = None
    ) -> None:
        target = (requirement_id or "").strip()
        if not target:
            return
        df = self.data_manager.dataframe
        if df.empty:
            self.log_console(f"No data loaded to locate: {target}")
            return
        lowered = target.lower()
        match_row: int | None = None
        match_source = None
        for idx, row in df.iterrows():
            for column in ("Requirement ID", "Up Trace"):
                value = str(row.get(column, "")).strip().lower()
                if value == lowered:
                    match_row = idx
                    match_source = str(row.get("SourceFile", "")).strip() or "Manual"
                    break
            if match_row is not None:
                break
        if match_row is None:
            self.log_console(f"Could not locate requirement: {target}")
            return
        row_source_type = (
            str(df.at[match_row, "SourceType"]).strip().lower()
            if "SourceType" in df.columns
            else ""
        )
        selected_source = source or match_source or "Manual"
        if row_source_type == "excel" and "Excel" in self._tab_source_types:
            selected_source = "Excel"
        self._select_source_tab(selected_source)
        self.view_stack.setCurrentWidget(self.table_tabs)
        self._select_global_row(selected_source, match_row)
        table = self._tab_tables.get(selected_source)
        if table is not None:
            table.setFocus()
        self.log_console(f"Navigated to requirement: {target}")

    # ------------------------------------------------------------------
    def show_trace_view(self) -> None:
        df = self.data_manager.to_trace_dataframe()
        if df.empty:
            QMessageBox.information(self, "Empty", "No data available.")
            return
        df = df.copy()
        df["_row_index"] = range(len(df))
        self.trace_view.load_data(df)
        self.view_stack.setCurrentIndex(2)
        self.log_console("Switched to Traceability Matrix view.")


class HeaderSettingsDialog(QDialog):
    """Modal dialog for editing document header, footer, and export options."""

    def __init__(self, parent: QWidget, settings: Dict[str, object]):
        super().__init__(parent)
        self.setWindowTitle("Header Details")
        self._settings = settings

        layout = QVBoxLayout(self)
        form = QFormLayout()

        self.title_edit = QLineEdit(str(settings.get("document_title", "")))
        form.addRow("Document Title", self.title_edit)

        self.doc_number_edit = QLineEdit(str(settings.get("document_number", "")))
        form.addRow("Document Number", self.doc_number_edit)

        self.revision_edit = QLineEdit(str(settings.get("revision", "")))
        form.addRow("Revision", self.revision_edit)

        self.author_name_edit = QLineEdit(str(settings.get("author_name", "")))
        form.addRow("Author Name", self.author_name_edit)
        self.author_title_edit = QLineEdit(str(settings.get("author_title", "")))
        form.addRow("Author Title", self.author_title_edit)

        self.reviewer_name_edit = QLineEdit(str(settings.get("reviewer_name", "")))
        form.addRow("Reviewer Name", self.reviewer_name_edit)
        self.reviewer_title_edit = QLineEdit(str(settings.get("reviewer_title", "")))
        form.addRow("Reviewer Title", self.reviewer_title_edit)

        self.qa_name_edit = QLineEdit(str(settings.get("qa_name", "")))
        form.addRow("QA Name", self.qa_name_edit)
        self.qa_title_edit = QLineEdit(str(settings.get("qa_title", "")))
        form.addRow("QA Title", self.qa_title_edit)

        self.config_name_edit = QLineEdit(str(settings.get("config_manager_name", "")))
        form.addRow("Configuration Manager Name", self.config_name_edit)
        self.config_title_edit = QLineEdit(
            str(settings.get("config_manager_title", ""))
        )
        form.addRow("Configuration Manager Title", self.config_title_edit)

        logo_left_widget, self.logo_left_edit = self._create_path_selector(
            str(settings.get("logo_left_path", ""))
        )
        form.addRow("Howell Logo", logo_left_widget)

        logo_right_widget, self.logo_right_edit = self._create_path_selector(
            str(settings.get("logo_right_path", ""))
        )
        form.addRow("Mach Logo", logo_right_widget)

        self.preview_spin = QSpinBox()
        self.preview_spin.setRange(10, 100)
        self.preview_spin.setSuffix(" %")
        self.preview_spin.setValue(int(settings.get("preview_image_width_percent", 80)))
        form.addRow("Preview Image Width", self.preview_spin)

        self.export_spin = QDoubleSpinBox()
        self.export_spin.setRange(1.0, 10.0)
        self.export_spin.setSingleStep(0.1)
        self.export_spin.setSuffix(' in')
        self.export_spin.setValue(float(settings.get("export_image_width_inches", 5.5)))
        form.addRow("Export Image Width", self.export_spin)

        self.watermark_edit = QLineEdit(str(settings.get("watermark_text", "")))
        form.addRow("Watermark Text", self.watermark_edit)

        self.address_howell_edit = QPlainTextEdit(
            str(settings.get("address_howell", "")).replace("<br/>", "\n")
        )
        form.addRow("Howell Address", self.address_howell_edit)

        self.address_mach_edit = QPlainTextEdit(
            str(settings.get("address_mach", "")).replace("<br/>", "\n")
        )
        form.addRow("Mach Address", self.address_mach_edit)

        self.notice_edit = QPlainTextEdit(
            str(settings.get("proprietary_notice", "")).replace("<br/>", "\n")
        )
        form.addRow("Proprietary Notice", self.notice_edit)

        self.copyright_edit = QLineEdit(str(settings.get("copyright_notice", "")))
        form.addRow("Copyright Notice", self.copyright_edit)

        layout.addLayout(form)

        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)

    def _create_path_selector(self, initial: str) -> tuple[QWidget, QLineEdit]:
        container = QWidget(self)
        h_layout = QHBoxLayout(container)
        h_layout.setContentsMargins(0, 0, 0, 0)
        line_edit = QLineEdit(initial)
        browse_btn = QPushButton("Browse")
        browse_btn.clicked.connect(lambda: self._browse_for_logo(line_edit))
        h_layout.addWidget(line_edit)
        h_layout.addWidget(browse_btn)
        return container, line_edit

    def _browse_for_logo(self, line_edit: QLineEdit) -> None:
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Logo",
            "",
            "Image Files (*.png *.jpg *.jpeg *.bmp *.gif *.tif *.tiff)",
        )
        if path:
            line_edit.setText(path)

    def values(self) -> Dict[str, object]:
        return {
            "document_title": self.title_edit.text().strip(),
            "document_number": self.doc_number_edit.text().strip(),
            "revision": self.revision_edit.text().strip(),
            "author_name": self.author_name_edit.text().strip(),
            "author_title": self.author_title_edit.text().strip(),
            "reviewer_name": self.reviewer_name_edit.text().strip(),
            "reviewer_title": self.reviewer_title_edit.text().strip(),
            "qa_name": self.qa_name_edit.text().strip(),
            "qa_title": self.qa_title_edit.text().strip(),
            "config_manager_name": self.config_name_edit.text().strip(),
            "config_manager_title": self.config_title_edit.text().strip(),
            "logo_left_path": self.logo_left_edit.text().strip(),
            "logo_right_path": self.logo_right_edit.text().strip(),
            "preview_image_width_percent": int(self.preview_spin.value()),
            "export_image_width_inches": float(self.export_spin.value()),
            "watermark_text": self.watermark_edit.text().strip(),
            "address_howell": self.address_howell_edit.toPlainText()
            .strip()
            .replace("\n", "<br/>"),
            "address_mach": self.address_mach_edit.toPlainText()
            .strip()
            .replace("\n", "<br/>"),
            "proprietary_notice": self.notice_edit.toPlainText()
            .strip()
            .replace("\n", "<br/>"),
            "copyright_notice": self.copyright_edit.text().strip(),
        }


def run_app() -> None:
    """Entry point used by scripts to launch the GUI."""
    import sys

    logging.basicConfig(level=logging.INFO)
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
# ------------------------------------------------------------------
if not hasattr(MainWindow, "show_trace_view"):
    def _fallback_show_trace_view(self) -> None:
        df = self.data_manager.to_trace_dataframe()
        if df.empty:
            QMessageBox.information(self, "Empty", "No data available.")
            return
        df = df.copy()
        df["_row_index"] = range(len(df))
        self.trace_view.load_data(df)
        self.view_stack.setCurrentIndex(2)
        self.log_console("Switched to Traceability Matrix view.")

    MainWindow.show_trace_view = _fallback_show_trace_view
